"""
Shared utilities for the RAG Proposal System.
DRY: functions used by multiple modules live here.
"""

import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path

from docx.shared import Pt, RGBColor


# ── Text Helpers ───────────────────────────────────────

def norm(s: str) -> str:
    """Normalize text for cheap dedup."""
    return re.sub(r"\s+", " ", s.lower()).strip()


def dedup_results(docs: List[str], metas: List[Dict]) -> Tuple[List[str], List[Dict]]:
    """Remove near-duplicate chunks."""
    seen: set = set()
    out_docs, out_metas = [], []
    for d, m in zip(docs, metas):
        key = norm(d[:900])
        if key not in seen:
            seen.add(key)
            out_docs.append(d)
            out_metas.append(m)
    return out_docs, out_metas


def format_citation(meta: Dict) -> str:
    """Format citation in human-friendly way: filename + page."""
    fn = meta.get("filename", "unknown")
    page = meta.get("page")
    sheet = meta.get("sheet")
    if sheet:
        return f"{fn} (sheet: {sheet})"
    if page:
        return f"{fn} p.{page}"
    return fn


def build_context(docs: List[str], metas: List[Dict], max_chars: int = 14000) -> str:
    """Build a context block with inline citations per chunk."""
    parts = []
    total = 0
    for i, (d, m) in enumerate(zip(docs, metas), start=1):
        cite = format_citation(m)
        block = f"[Source {i}: {cite}]\n{d.strip()}\n"
        if total + len(block) > max_chars:
            break
        parts.append(block)
        total += len(block)
    return "\n".join(parts)


# ── Opportunity Detection ──────────────────────────────

def detect_opportunity_from_query(query: str, known_opportunities: List[str]) -> Optional[str]:
    """
    Auto-detect opportunity ID from query text.
    Checks explicit IDs (CORHQ-25-R-0450) and fuzzy name matches.
    """
    # Explicit ID pattern
    match = re.search(r'[A-Z]+[-_]\d{2}[-_][A-Z][-_]\d{4}', query, re.IGNORECASE)
    if match:
        opp_id = match.group(0).upper()
        for opp in known_opportunities:
            if opp_id in opp.upper():
                return opp

    # Fuzzy match
    query_lower = query.lower()
    for opp in known_opportunities:
        opp_normalized = opp.lower().replace('_', ' ').replace('-', ' ')
        if opp_normalized in query_lower or opp.lower() in query_lower:
            return opp

    return None


# ── DOCX Formatting ────────────────────────────────────

def add_formatted_text(paragraph, text: str, font_name: str = 'Times New Roman',
                       font_size_pt: int = 12, color_rgb: tuple = (0, 0, 0)):
    """
    Add text to a python-docx paragraph, converting **bold** and *italic* markdown.
    """
    pattern = r'(\*\*.*?\*\*|\*.*?\*|[^\*]+|\*)'
    parts = re.findall(pattern, text)

    for part in parts:
        if not part or part == '*':
            continue

        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            run = paragraph.add_run(part)

        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.font.color.rgb = RGBColor(*color_rgb)


# ── ChromaDB Helpers ───────────────────────────────────

def get_chroma_client(db_path: str = None):
    """Get a ChromaDB persistent client."""
    import chromadb
    from config import CHROMA_DB_PATH
    path = db_path or str(CHROMA_DB_PATH)
    return chromadb.PersistentClient(path=path)


def get_embedder(api_key: str = None):
    """Get OpenAI embedding function for ChromaDB."""
    import os
    from chromadb.utils import embedding_functions
    from config import EMBEDDING_MODEL
    key = api_key or os.environ["OPENAI_API_KEY"]
    return embedding_functions.OpenAIEmbeddingFunction(
        api_key=key,
        model_name=EMBEDDING_MODEL,
    )
