"""
Requirement Tracker
Tracks which RFP requirements are addressed in proposal sections.

Analyzes proposal drafts against extracted requirements to show coverage gaps.
"""

import os
import re
from typing import List, Dict, Tuple
from pathlib import Path
from datetime import datetime

import chromadb
from chromadb.utils import embedding_functions
from docx import Document
from colorama import init, Fore, Style

init(autoreset=True)


def load_requirements_from_rfp(opportunity: str, coll) -> List[Dict]:
    """Extract requirements from RFP documents."""
    where = {
        "$and": [
            {"opportunity": opportunity},
            {"stage": {"$ne": "award"}},
            {"doc_role": {"$in": ["technical_requirements", "instructions", "evaluation_criteria"]}}
        ]
    }
    
    query_result = coll.query(
        query_texts=["requirements specifications shall must will should"],
        n_results=30,
        where=where
    )
    
    requirements = []
    docs = query_result["documents"][0]
    metas = query_result["metadatas"][0]
    
    req_id = 1
    for doc, meta in zip(docs, metas):
        sentences = re.split(r'(?<=[.!?])\s+', doc)
        
        for sentence in sentences:
            if len(sentence) < 20:
                continue
            
            sentence = sentence.strip()
            
            # Identify requirement type
            if re.search(r'\b(shall|must)\b', sentence, re.IGNORECASE):
                requirements.append({
                    'id': f"REQ-{req_id:04d}",
                    'text': sentence[:300],
                    'source': meta.get('filename', 'Unknown'),
                    'page': meta.get('page'),
                    'addressed': False,
                    'proposal_locations': []
                })
                req_id += 1
    
    return requirements


def load_proposal_content(proposal_path: Path) -> Dict[str, str]:
    """Load proposal sections from Word documents or text files."""
    sections = {}
    
    if proposal_path.is_dir():
        # Load all .docx and .txt files
        for file_path in proposal_path.glob("**/*"):
            if file_path.suffix in ['.docx', '.txt']:
                section_name = file_path.stem
                
                if file_path.suffix == '.docx':
                    try:
                        doc = Document(str(file_path))
                        content = "\n".join([p.text for p in doc.paragraphs])
                        sections[section_name] = content
                    except Exception as e:
                        print(f"Warning: Could not read {file_path}: {e}")
                elif file_path.suffix == '.txt':
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            sections[section_name] = f.read()
                    except Exception as e:
                        print(f"Warning: Could not read {file_path}: {e}")
    elif proposal_path.is_file():
        # Single file
        section_name = proposal_path.stem
        
        if proposal_path.suffix == '.docx':
            doc = Document(str(proposal_path))
            content = "\n".join([p.text for p in doc.paragraphs])
            sections[section_name] = content
        elif proposal_path.suffix == '.txt':
            with open(proposal_path, 'r', encoding='utf-8') as f:
                sections[section_name] = f.read()
    
    return sections


def check_requirement_coverage(requirement: Dict, sections: Dict[str, str]) -> Tuple[bool, List[str]]:
    """
    Check if a requirement is addressed in proposal sections.
    
    Returns (is_addressed, list_of_sections)
    """
    req_text_lower = requirement['text'].lower()
    
    # Extract key terms (nouns and technical terms)
    key_terms = []
    words = re.findall(r'\b[a-z]{4,}\b', req_text_lower)
    
    # Filter out common words
    common_words = {'shall', 'must', 'will', 'should', 'provide', 'include', 'ensure', 
                    'system', 'requirements', 'contractor', 'government'}
    key_terms = [w for w in words if w not in common_words][:5]  # Top 5 key terms
    
    if not key_terms:
        return False, []
    
    locations = []
    for section_name, section_content in sections.items():
        section_lower = section_content.lower()
        
        # Check if multiple key terms appear
        matches = sum(1 for term in key_terms if term in section_lower)
        
        if matches >= min(2, len(key_terms)):  # At least 2 terms or all if less than 2
            locations.append(section_name)
    
    is_addressed = len(locations) > 0
    return is_addressed, locations


def generate_coverage_report(requirements: List[Dict], sections: Dict[str, str]) -> str:
    """Generate text report of requirement coverage."""
    addressed = [r for r in requirements if r['addressed']]
    not_addressed = [r for r in requirements if not r['addressed']]
    
    coverage_pct = (len(addressed) / len(requirements) * 100) if requirements else 0
    
    report = []
    report.append("=" * 80)
    report.append("REQUIREMENT COVERAGE REPORT")
    report.append("=" * 80)
    report.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append("")
    report.append(f"Total Requirements: {len(requirements)}")
    report.append(f"Addressed: {len(addressed)} ({coverage_pct:.1f}%)")
    report.append(f"Not Addressed: {len(not_addressed)}")
    report.append("")
    
    if not_addressed:
        report.append("=" * 80)
        report.append("GAPS - Requirements NOT Addressed:")
        report.append("=" * 80)
        
        for req in not_addressed:
            report.append(f"\n{req['id']} [{req['source']}]")
            report.append(f"  {req['text'][:200]}")
    
    if addressed:
        report.append("\n" + "=" * 80)
        report.append("COVERAGE - Requirements Addressed:")
        report.append("=" * 80)
        
        for req in addressed:
            report.append(f"\n{req['id']} [{req['source']}]")
            report.append(f"  Location(s): {', '.join(req['proposal_locations'])}")
            report.append(f"  {req['text'][:150]}...")
    
    return "\n".join(report)


def main():
    if not os.getenv("OPENAI_API_KEY"):
        raise SystemExit("OPENAI_API_KEY not set.\n")
    
    print("\n" + "=" * 60)
    print("REQUIREMENT TRACKER")
    print("=" * 60)
    
    # Setup
    client_chroma = chromadb.PersistentClient(path="chroma_db")
    embedder = embedding_functions.OpenAIEmbeddingFunction(
        api_key=os.environ["OPENAI_API_KEY"],
        model_name="text-embedding-3-large"
    )
    coll = client_chroma.get_collection("authoritative", embedding_function=embedder)
    
    # Get available opportunities
    result = coll.get()
    opportunities = sorted(set(m.get('opportunity', 'unknown') for m in result['metadatas'] if m.get('opportunity') != 'unknown'))
    
    print("\nAvailable opportunities:")
    for i, opp in enumerate(opportunities, 1):
        print(f"  {i}. {opp}")
    
    choice = input("\nSelect opportunity (1-N): ").strip()
    opportunity = opportunities[int(choice) - 1]
    
    print(f"\n{Fore.CYAN}→ Loading requirements for {opportunity}...{Style.RESET_ALL}")
    requirements = load_requirements_from_rfp(opportunity, coll)
    print(f"  Found {len(requirements)} requirements")
    
    # Get proposal content to check
    proposal_input = input("\nEnter path to proposal draft (file or folder): ").strip()
    proposal_path = Path(proposal_input)
    
    if not proposal_path.exists():
        print(f"{Fore.RED}Error: Path does not exist{Style.RESET_ALL}")
        return
    
    print(f"\n{Fore.CYAN}→ Loading proposal content...{Style.RESET_ALL}")
    sections = load_proposal_content(proposal_path)
    print(f"  Loaded {len(sections)} section(s)")
    
    for section_name in sections.keys():
        print(f"    - {section_name}")
    
    # Check coverage
    print(f"\n{Fore.CYAN}→ Analyzing requirement coverage...{Style.RESET_ALL}")
    
    for req in requirements:
        is_addressed, locations = check_requirement_coverage(req, sections)
        req['addressed'] = is_addressed
        req['proposal_locations'] = locations
    
    # Generate report
    report = generate_coverage_report(requirements, sections)
    
    # Display summary
    addressed_count = sum(1 for r in requirements if r['addressed'])
    coverage_pct = (addressed_count / len(requirements) * 100) if requirements else 0
    
    print("\n" + "=" * 60)
    print("COVERAGE SUMMARY")
    print("=" * 60)
    print(f"Total Requirements: {len(requirements)}")
    print(f"{Fore.GREEN}Addressed: {addressed_count} ({coverage_pct:.1f}%){Style.RESET_ALL}")
    print(f"{Fore.RED}Not Addressed: {len(requirements) - addressed_count}{Style.RESET_ALL}")
    
    if coverage_pct < 80:
        print(f"\n{Fore.YELLOW}⚠️  WARNING: Coverage below 80%. Review gaps before submission.{Style.RESET_ALL}")
    elif coverage_pct < 100:
        print(f"\n{Fore.YELLOW}⚠️  Good progress, but some requirements still need attention.{Style.RESET_ALL}")
    else:
        print(f"\n{Fore.GREEN}✅ All requirements addressed!{Style.RESET_ALL}")
    
    # Save report
    report_file = f"{opportunity}_Coverage_Report.txt"
    with open(report_file, 'w', encoding='utf-8') as f:
        f.write(report)
    
    print(f"\n📄 Detailed report saved: {report_file}")
    
    # Show top gaps
    not_addressed = [r for r in requirements if not r['addressed']]
    if not_addressed:
        print(f"\n{Fore.RED}Top 5 gaps to address:{Style.RESET_ALL}")
        for req in not_addressed[:5]:
            print(f"\n  {req['id']} - {req['source']}")
            print(f"    {req['text'][:120]}...")


if __name__ == "__main__":
    main()
