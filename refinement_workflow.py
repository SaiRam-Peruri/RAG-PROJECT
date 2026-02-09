"""
Multi-Round Refinement Workflow
Iteratively improves proposal sections based on feedback and quality checks.

Process:
1. Generate initial draft
2. Analyze for issues (compliance, completeness, quality)
3. Apply feedback and regenerate
4. Track versions and changes
5. Repeat until quality threshold met
"""

import os
import re
from typing import List, Dict, Tuple, Optional
from pathlib import Path
from datetime import datetime
import difflib

import chromadb
from chromadb.utils import embedding_functions
from openai import OpenAI
from docx import Document
from docx.shared import RGBColor, Pt
from colorama import init, Fore, Style

init(autoreset=True)


class SectionRefinementEngine:
    """Manages iterative refinement of proposal sections."""
    
    def __init__(self, opportunity: str, section_type: str):
        self.opportunity = opportunity
        self.section_type = section_type
        self.versions = []  # List of (version_num, content, feedback, improvements)
        
        self.client_openai = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
        self.client_chroma = chromadb.PersistentClient(path="chroma_db")
        embedder = embedding_functions.OpenAIEmbeddingFunction(
            api_key=os.environ["OPENAI_API_KEY"],
            model_name="text-embedding-3-large"
        )
        self.coll_auth = self.client_chroma.get_collection("authoritative", embedding_function=embedder)
        self.coll_draft = self.client_chroma.get_collection("drafting", embedding_function=embedder)
    
    def generate_initial_draft(self) -> str:
        """Generate version 1 draft."""
        print(f"{Fore.CYAN}→ Generating initial draft (v1)...{Style.RESET_ALL}")
        
        # Query requirements
        requirements_context = self._query_requirements()
        
        # Query best practices
        best_practices_context = self._query_best_practices()
        
        # Generate
        content = self._generate_content(requirements_context, best_practices_context, version=1)
        
        self.versions.append({
            'version': 1,
            'content': content,
            'feedback': [],
            'quality_score': self._calculate_quality_score(content, requirements_context)
        })
        
        return content
    
    def analyze_draft(self, content: str) -> Dict:
        """
        Analyze draft for issues.
        
        Returns dict with:
        - compliance_check: List of missing requirements
        - quality_check: Readability, clarity, structure issues
        - completeness_check: Missing sections or details
        - recommendations: Specific improvements
        """
        print(f"{Fore.CYAN}→ Analyzing draft quality...{Style.RESET_ALL}")
        
        analysis_prompt = f"""You are a federal proposal quality reviewer. Analyze this {self.section_type} section draft.

DRAFT:
{content[:4000]}

Provide analysis in this format:

COMPLIANCE:
- [List any missing mandatory requirements or evaluation factors]

QUALITY:
- [Issues with clarity, structure, readability, persuasiveness]

COMPLETENESS:
- [Missing details, examples, or supporting evidence]

RECOMMENDATIONS:
- [Top 3-5 specific improvements, prioritized]

Be specific and actionable. Focus on high-impact improvements."""

        response = self.client_openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": analysis_prompt}],
            temperature=0.3
        )
        
        analysis_text = response.choices[0].message.content
        
        # Parse analysis into structured format
        analysis = {
            'compliance': self._extract_section(analysis_text, 'COMPLIANCE'),
            'quality': self._extract_section(analysis_text, 'QUALITY'),
            'completeness': self._extract_section(analysis_text, 'COMPLETENESS'),
            'recommendations': self._extract_section(analysis_text, 'RECOMMENDATIONS'),
            'full_text': analysis_text
        }
        
        return analysis
    
    def refine_draft(self, feedback: str) -> str:
        """Generate improved version based on feedback."""
        current_version = self.versions[-1]
        new_version_num = current_version['version'] + 1
        
        print(f"{Fore.CYAN}→ Generating refined draft (v{new_version_num})...{Style.RESET_ALL}")
        
        # Get context again
        requirements_context = self._query_requirements()
        best_practices_context = self._query_best_practices()
        
        # Generate improved version
        refinement_prompt = f"""You are refining a federal proposal section based on feedback.

SECTION TYPE: {self.section_type}
OPPORTUNITY: {self.opportunity}

PREVIOUS VERSION (v{current_version['version']}):
{current_version['content'][:3000]}

FEEDBACK TO ADDRESS:
{feedback}

GOVERNMENT REQUIREMENTS:
{requirements_context[:3000]}

INTERNAL BEST PRACTICES:
{best_practices_context[:2000]}

Generate an improved version that:
1. Addresses ALL feedback points
2. Maintains what was working well
3. Improves clarity and persuasiveness
4. Ensures full compliance with requirements

Output the improved section directly (no meta-commentary)."""

        response = self.client_openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": refinement_prompt}],
            temperature=0.4
        )
        
        improved_content = response.choices[0].message.content.strip()
        
        # Track changes
        changes = self._identify_changes(current_version['content'], improved_content)
        
        self.versions.append({
            'version': new_version_num,
            'content': improved_content,
            'feedback': feedback,
            'changes': changes,
            'quality_score': self._calculate_quality_score(improved_content, requirements_context)
        })
        
        return improved_content
    
    def _query_requirements(self) -> str:
        """Query government requirements."""
        where = {
            "$and": [
                {"opportunity": self.opportunity},
                {"authority": "government"},
                {"doc_role": {"$in": ["technical_requirements", "evaluation_criteria", "instructions"]}}
            ]
        }
        
        section_queries = {
            'technical': 'technical approach architecture solution system design',
            'management': 'management plan organization quality assurance staffing',
            'past_performance': 'past performance relevant experience similar contracts',
            'executive_summary': 'executive summary mission objectives capability'
        }
        
        query_text = section_queries.get(self.section_type, f'{self.section_type} requirements')
        
        result = self.coll_auth.query(
            query_texts=[query_text],
            n_results=10,
            where=where
        )
        
        return "\n\n".join(result["documents"][0][:5])
    
    def _query_best_practices(self) -> str:
        """Query internal best practices."""
        where = {
            "authority": "vendor"
        }
        
        result = self.coll_draft.query(
            query_texts=[f"{self.section_type} best practices approach strategy"],
            n_results=5,
            where=where
        )
        
        if result["documents"][0]:
            return "\n\n".join(result["documents"][0][:3])
        return ""
    
    def _generate_content(self, requirements: str, best_practices: str, version: int) -> str:
        """Generate section content."""
        section_prompts = {
            'technical': "Address technical requirements comprehensively. Structure: (1) Understanding of Requirements, (2) Proposed Solution, (3) Technical Implementation Details.",
            'management': "Address organizational structure, key personnel roles, quality assurance processes, and risk management. Show proven processes.",
            'past_performance': "Reference specific contracts demonstrating capability. Include: customer, contract value, dates, relevance to current RFP, outcomes/results.",
            'executive_summary': "Synthesize: (1) Understanding of agency mission, (2) Unique advantages/differentiators, (3) Why we are the best choice."
        }
        
        prompt = section_prompts.get(self.section_type, f"Draft a {self.section_type} section.")
        
        generate_prompt = f"""You are drafting a federal proposal section.

SECTION: {self.section_type}
OPPORTUNITY: {self.opportunity}
VERSION: {version}

GOVERNMENT REQUIREMENTS:
{requirements[:3000]}

INTERNAL BEST PRACTICES:
{best_practices[:2000]}

INSTRUCTIONS:
{prompt}

Write a complete, professional section. Be specific, use active voice, demonstrate capability."""

        response = self.client_openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": generate_prompt}],
            temperature=0.4
        )
        
        return response.choices[0].message.content.strip()
    
    def _extract_section(self, text: str, section_name: str) -> List[str]:
        """Extract bullet points from analysis section."""
        pattern = f"{section_name}:(.*?)(?=\n[A-Z]+:|$)"
        match = re.search(pattern, text, re.DOTALL)
        
        if match:
            section_text = match.group(1)
            bullets = [line.strip() for line in section_text.split('\n') if line.strip().startswith('-')]
            return bullets
        
        return []
    
    def _identify_changes(self, old: str, new: str) -> List[str]:
        """Identify major changes between versions."""
        differ = difflib.Differ()
        old_lines = old.split('\n')
        new_lines = new.split('\n')
        
        diff = list(differ.compare(old_lines, new_lines))
        
        changes = []
        added = [line[2:] for line in diff if line.startswith('+ ')]
        removed = [line[2:] for line in diff if line.startswith('- ')]
        
        if added:
            changes.append(f"Added {len(added)} new lines/paragraphs")
        if removed:
            changes.append(f"Removed {len(removed)} lines")
        
        return changes
    
    def _calculate_quality_score(self, content: str, requirements: str) -> float:
        """Calculate quality score 0-100."""
        score = 70  # Base score
        
        # Length check
        if len(content) > 1000:
            score += 5
        if len(content) > 2500:
            score += 5
        
        # Structure markers
        if any(marker in content for marker in ['Understanding', 'Approach', 'Implementation']):
            score += 5
        
        # Specificity (numbers, examples)
        if len(re.findall(r'\d+', content)) > 5:
            score += 5
        
        # Active voice indicators
        active_verbs = ['will', 'provide', 'deliver', 'implement', 'ensure', 'develop']
        if sum(content.lower().count(verb) for verb in active_verbs) > 10:
            score += 5
        
        # Requirement alignment (simple keyword overlap)
        req_keywords = set(re.findall(r'\b[a-z]{5,}\b', requirements.lower()))
        content_keywords = set(re.findall(r'\b[a-z]{5,}\b', content.lower()))
        overlap = len(req_keywords & content_keywords)
        
        if overlap > 10:
            score += 5
        
        return min(score, 100)
    
    def save_version_history(self, output_path: Path):
        """Save all versions with tracked changes to Word document."""
        doc = Document()
        
        doc.add_heading(f"{self.opportunity} - {self.section_type.title()} Version History", 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Versions: {len(self.versions)}")
        
        for ver in self.versions:
            doc.add_page_break()
            
            heading = doc.add_heading(f"Version {ver['version']}", 1)
            
            # Quality score
            score_para = doc.add_paragraph(f"Quality Score: {ver['quality_score']:.1f}/100")
            if ver['quality_score'] >= 85:
                score_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif ver['quality_score'] >= 70:
                score_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)  # Orange
            else:
                score_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            # Feedback applied (if not first version)
            if ver['version'] > 1:
                doc.add_heading("Feedback Addressed:", 2)
                doc.add_paragraph(ver['feedback'])
                
                if 'changes' in ver:
                    doc.add_heading("Changes Made:", 2)
                    for change in ver['changes']:
                        doc.add_paragraph(change, style='List Bullet')
            
            # Content
            doc.add_heading("Content:", 2)
            doc.add_paragraph(ver['content'])
        
        doc.save(str(output_path))


def main():
    if not os.getenv("OPENAI_API_KEY"):
        raise SystemExit("OPENAI_API_KEY not set.\n")
    
    print("\n" + "=" * 70)
    print("MULTI-ROUND REFINEMENT WORKFLOW")
    print("=" * 70)
    
    # Get inputs
    opportunity = input("\nEnter opportunity name (e.g., CORHQ-25-R-0450): ").strip()
    
    print("\nSection types:")
    print("  1. Technical Approach")
    print("  2. Management Plan")
    print("  3. Past Performance")
    print("  4. Executive Summary")
    
    section_choice = input("\nSelect section (1-4): ").strip()
    section_map = {'1': 'technical', '2': 'management', '3': 'past_performance', '4': 'executive_summary'}
    section_type = section_map.get(section_choice, 'technical')
    
    # Initialize engine
    engine = SectionRefinementEngine(opportunity, section_type)
    
    # Round 1: Initial draft
    print(f"\n{Fore.GREEN}{'='*70}{Style.RESET_ALL}")
    print(f"{Fore.GREEN}ROUND 1: Initial Draft{Style.RESET_ALL}")
    print(f"{Fore.GREEN}{'='*70}{Style.RESET_ALL}")
    
    draft_v1 = engine.generate_initial_draft()
    v1_score = engine.versions[-1]['quality_score']
    
    print(f"\n✅ Draft v1 generated (Quality: {v1_score:.1f}/100)")
    print(f"\nPreview (first 500 chars):\n{draft_v1[:500]}...\n")
    
    # Analyze
    analysis = engine.analyze_draft(draft_v1)
    
    print(f"\n{Fore.YELLOW}ANALYSIS RESULTS:{Style.RESET_ALL}")
    for category, items in [('COMPLIANCE', analysis['compliance']), 
                             ('QUALITY', analysis['quality']),
                             ('COMPLETENESS', analysis['completeness'])]:
        if items:
            print(f"\n{category}:")
            for item in items[:3]:  # Show top 3
                print(f"  {item}")
    
    # Refinement rounds
    round_num = 2
    max_rounds = 5
    
    while round_num <= max_rounds:
        proceed = input(f"\n{Fore.CYAN}Proceed with refinement round {round_num}? (y/n/view): {Style.RESET_ALL}").strip().lower()
        
        if proceed == 'n':
            break
        elif proceed == 'view':
            print(f"\n{analysis['full_text']}")
            continue
        
        print(f"\n{Fore.GREEN}{'='*70}{Style.RESET_ALL}")
        print(f"{Fore.GREEN}ROUND {round_num}: Refinement{Style.RESET_ALL}")
        print(f"{Fore.GREEN}{'='*70}{Style.RESET_ALL}")
        
        # Get feedback (could be manual or use analysis)
        use_auto = input("Use automated feedback from analysis? (y/n): ").strip().lower() == 'y'
        
        if use_auto:
            feedback = "\n".join([
                "FOCUS AREAS:",
                *analysis['recommendations'][:5]
            ])
        else:
            print("\nEnter feedback (type END on new line when done):")
            feedback_lines = []
            while True:
                line = input()
                if line == "END":
                    break
                feedback_lines.append(line)
            feedback = "\n".join(feedback_lines)
        
        # Refine
        improved_draft = engine.refine_draft(feedback)
        improved_score = engine.versions[-1]['quality_score']
        
        print(f"\n✅ Draft v{round_num} generated (Quality: {improved_score:.1f}/100)")
        print(f"   Improvement: {improved_score - v1_score:+.1f} points")
        
        # Re-analyze if continuing
        if round_num < max_rounds:
            analysis = engine.analyze_draft(improved_draft)
        
        round_num += 1
    
    # Save final output
    final_version = engine.versions[-1]
    final_content = final_version['content']
    final_score = final_version['quality_score']
    
    output_file = Path(f"{opportunity}_{section_type}_REFINED_v{final_version['version']}.docx")
    
    doc = Document()
    doc.add_heading(f"{section_type.replace('_', ' ').title()}", 0)
    doc.add_paragraph(f"Opportunity: {opportunity}")
    doc.add_paragraph(f"Version: {final_version['version']}")
    doc.add_paragraph(f"Quality Score: {final_score:.1f}/100")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("")
    doc.add_paragraph(final_content)
    doc.save(str(output_file))
    
    # Save version history
    history_file = Path(f"{opportunity}_{section_type}_VERSION_HISTORY.docx")
    engine.save_version_history(history_file)
    
    print(f"\n{Fore.GREEN}{'='*70}{Style.RESET_ALL}")
    print(f"{Fore.GREEN}REFINEMENT COMPLETE{Style.RESET_ALL}")
    print(f"{Fore.GREEN}{'='*70}{Style.RESET_ALL}")
    print(f"\nFinal Version: v{final_version['version']}")
    print(f"Final Quality Score: {final_score:.1f}/100")
    print(f"\n📄 Output saved: {output_file}")
    print(f"📚 Version history: {history_file}")
    
    if final_score >= 85:
        print(f"\n{Fore.GREEN}✅ High quality - Ready for review{Style.RESET_ALL}")
    elif final_score >= 70:
        print(f"\n{Fore.YELLOW}⚠️  Good quality - Consider one more refinement pass{Style.RESET_ALL}")
    else:
        print(f"\n{Fore.RED}⚠️  Needs improvement - Additional refinement recommended{Style.RESET_ALL}")


if __name__ == "__main__":
    main()
