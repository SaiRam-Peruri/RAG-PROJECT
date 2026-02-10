"""
Proposal Section Generator
Auto-drafts proposal sections using authoritative + drafting collections.

This is how companies actually use AI for proposals:
- Section-by-section generation
- Requirements + internal best practices
- Full citations
"""

import os
import re
from typing import Dict, List, Tuple, Optional
from pathlib import Path

import chromadb
from chromadb.utils import embedding_functions
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def detect_opportunity_from_query(query: str, all_opportunities: List[str]) -> Optional[str]:
    """
    Auto-detect opportunity ID from query text.
    
    Patterns:
    - CORHQ-25-R-0450 (explicit ID)
    - "DMS Support" → DMS_Support
    - Fuzzy match against known opportunities
    """
    # Pattern 1: Explicit opportunity ID (CORHQ-XX-X-XXXX format)
    pattern = r'[A-Z]+[-_]\d{2}[-_][A-Z][-_]\d{4}'
    match = re.search(pattern, query, re.IGNORECASE)
    if match:
        return match.group(0).upper()
    
    # Pattern 2: Match against known opportunity names
    query_lower = query.lower()
    for opp in all_opportunities:
        # Normalize opportunity name for matching
        opp_normalized = opp.lower().replace('_', ' ').replace('-', ' ')
        if opp_normalized in query_lower or opp.lower() in query_lower:
            return opp
    
    return None


def get_available_opportunities(client_chroma) -> List[str]:
    """Get list of all opportunities in the database."""
    coll = client_chroma.get_collection("authoritative")
    result = coll.get()
    opportunities = set(m.get('opportunity', 'unknown') for m in result['metadatas'])
    opportunities.discard('unknown')
    return sorted(opportunities)


def query_requirements(coll, opportunity: str, section_focus: str = None) -> Tuple[List[str], List[Dict]]:
    """
    Retrieve requirements for a specific section.
    
    Args:
        section_focus: 'technical', 'management', 'past_performance', etc.
    """
    query_text = f"requirements specifications {section_focus or 'technical'}"
    
    where = {
        "$and": [
            {"stage": {"$ne": "award"}},
            {"opportunity": opportunity},
            {"doc_role": {"$in": ["technical_requirements", "evaluation_criteria", "instructions", "amendment_qa"]}}
        ]
    }
    
    result = coll.query(
        query_texts=[query_text],
        n_results=8,
        where=where
    )
    
    return result["documents"][0], result["metadatas"][0]


def query_past_wins(coll_draft, section_type: str) -> Tuple[List[str], List[Dict]]:
    """
    Retrieve internal best practices from past proposals.
    
    section_type: 'technical', 'management', 'past_performance'
    """
    query_text = f"{section_type} approach best practices"
    
    where = {
        "$and": [
            {"authority": "vendor"},
            {"doc_role": {"$in": [section_type, "proposal"]}}
        ]
    }
    
    result = coll_draft.query(
        query_texts=[query_text],
        n_results=5,
        where=where
    )
    
    return result["documents"][0], result["metadatas"][0]


def build_section_context(req_docs: List[str], req_metas: List[Dict], 
                          past_docs: List[str] = None, past_metas: List[Dict] = None) -> str:
    """Build context with clear separation of requirements vs. best practices."""
    
    context_parts = ["=== GOVERNMENT REQUIREMENTS ===\n"]
    
    for i, (doc, meta) in enumerate(zip(req_docs, req_metas), 1):
        cite = f"{meta.get('filename', 'unknown')} p.{meta.get('page', '?')}"
        context_parts.append(f"[Requirement {i}: {cite}]\n{doc.strip()}\n")
    
    if past_docs and past_metas:
        context_parts.append("\n=== INTERNAL BEST PRACTICES (Past Wins) ===\n")
        for i, (doc, meta) in enumerate(zip(past_docs, past_metas), 1):
            cite = f"{meta.get('filename', 'unknown')}"
            context_parts.append(f"[Best Practice {i}: {cite}]\n{doc.strip()}\n")
    
    return "\n".join(context_parts)


def generate_section(client_openai: OpenAI, section_type: str, opportunity: str, 
                     context: str, requirements_metas: List[Dict]) -> Dict:
    """
    Generate a proposal section draft.
    
    Returns:
        {
            'content': str,
            'requirements_addressed': List[str],
            'citations': List[str]
        }
    """
    
    section_prompts = {
        'technical': (
            "You are writing the Technical Approach section of a federal proposal. "
            "Address ALL requirements from the GOVERNMENT REQUIREMENTS section. "
            "Use INTERNAL BEST PRACTICES as examples but adapt to current requirements. "
            "Structure: (1) Understanding of Requirements, (2) Technical Solution, (3) Implementation Approach. "
            "Cite requirements with [Requirement N]. "
            "Be specific, avoid generic statements."
        ),
        'management': (
            "You are writing the Management Plan section of a federal proposal. "
            "Address organizational structure, key personnel roles, quality assurance, and risk management. "
            "Use GOVERNMENT REQUIREMENTS to understand what's needed. "
            "Use INTERNAL BEST PRACTICES as proven frameworks. "
            "Cite requirements with [Requirement N]."
        ),
        'past_performance': (
            "You are writing the Past Performance section of a federal proposal. "
            "Reference specific past contracts that demonstrate relevant experience. "
            "Use GOVERNMENT REQUIREMENTS to understand evaluation criteria. "
            "For each example, include: customer, contract value, dates, relevance, outcomes. "
            "Cite requirements with [Requirement N]."
        ),
        'executive_summary': (
            "You are writing the Executive Summary of a federal proposal. "
            "Synthesize: (1) Understanding of mission, (2) Unique solution advantages, (3) Why we win. "
            "Keep it concise (1-2 pages). "
            "Use GOVERNMENT REQUIREMENTS to show alignment. "
            "Cite requirements with [Requirement N]."
        ),
        'staffing': (
            "You are writing the Staffing Plan section of a federal proposal. "
            "Address: (1) Key personnel qualifications, (2) Org chart, (3) Labor categories, (4) Recruitment strategy. "
            "Use GOVERNMENT REQUIREMENTS for required labor categories and qualifications. "
            "Use INTERNAL BEST PRACTICES for proven staffing models. "
            "Include specific roles, responsibilities, and qualifications. "
            "Cite requirements with [Requirement N]."
        ),
        'quality_assurance': (
            "You are writing the Quality Assurance section of a federal proposal. "
            "Address: (1) QA/QC processes, (2) Quality metrics, (3) Testing procedures, (4) Continuous improvement. "
            "Use GOVERNMENT REQUIREMENTS for quality standards and acceptance criteria. "
            "Use INTERNAL BEST PRACTICES for proven QA frameworks. "
            "Cite requirements with [Requirement N]."
        ),
        'security': (
            "You are writing the Security and Compliance section of a federal proposal. "
            "Address: (1) Security controls, (2) Compliance frameworks (FedRAMP, FISMA, NIST), (3) Data protection, (4) Access management. "
            "Use GOVERNMENT REQUIREMENTS for security and compliance mandates. "
            "Use INTERNAL BEST PRACTICES for proven security implementations. "
            "Cite requirements with [Requirement N]."
        ),
        'transition': (
            "You are writing the Transition Plan section of a federal proposal. "
            "Address: (1) Transition approach, (2) Knowledge transfer, (3) Risk mitigation, (4) Timeline and milestones. "
            "Use GOVERNMENT REQUIREMENTS for transition requirements and constraints. "
            "Use INTERNAL BEST PRACTICES for proven transition methodologies. "
            "Cite requirements with [Requirement N]."
        ),
        'cost': (
            "You are writing the Cost Proposal section of a federal proposal. "
            "Address: (1) Cost breakdown, (2) Pricing strategy, (3) Cost justification, (4) Value proposition. "
            "Use GOVERNMENT REQUIREMENTS for pricing format and evaluation criteria. "
            "Be specific about labor rates, ODCs, and total cost. "
            "Cite requirements with [Requirement N]."
        )
    }
    
    system = section_prompts.get(section_type, section_prompts['technical'])
    
    user = (
        f"OPPORTUNITY: {opportunity}\n"
        f"SECTION: {section_type.replace('_', ' ').title()}\n\n"
        f"CONTEXT:\n{context}\n\n"
        "INSTRUCTIONS:\n"
        "- Address every requirement explicitly\n"
        "- Use [Requirement N] citations\n"
        "- Be specific and detailed (not generic)\n"
        "- Show understanding + solution + benefits\n"
        "- Draft 2-3 pages of content\n"
    )
    
    response = client_openai.chat.completions.create(
        model="gpt-4o-mini",  # Fast for drafting
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user}
        ],
        temperature=0.3,  # Some creativity, but controlled
        max_tokens=3000
    )
    
    content = response.choices[0].message.content
    
    return {
        'content': content,
        'citations': [f"{m.get('filename')} p.{m.get('page')}" for m in requirements_metas]
    }


def create_docx_output(section_type: str, opportunity: str, content: str, 
                       citations: List[str], output_path: Path):
    """Create a Word document with the generated section."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f"{section_type.replace('_', ' ').title()}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Opportunity: {opportunity}\n").bold = True
    meta.add_run(f"Generated: {Path(__file__).name}\n")
    
    doc.add_page_break()
    
    # Content
    doc.add_heading('Draft Content', 1)
    
    # Split content into paragraphs
    for para in content.split('\n\n'):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.style = 'Normal'
    
    doc.add_page_break()
    
    # Citations
    doc.add_heading('Source Documents', 1)
    doc.add_paragraph("This draft was generated using the following sources:")
    
    for i, cite in enumerate(citations, 1):
        doc.add_paragraph(f"{i}. {cite}", style='List Number')
    
    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph(
        "NOTE: This is an AI-generated draft. Review and edit for accuracy, "
        "completeness, and alignment with capture strategy before submission."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(10)
    
    doc.save(str(output_path))


def main():
    if not os.getenv("OPENAI_API_KEY"):
        raise SystemExit("OPENAI_API_KEY not set.\n")
    
    print("\n=== Proposal Section Generator ===\n")
    
    # Setup
    client_chroma = chromadb.PersistentClient(path="chroma_db")
    embedder = embedding_functions.OpenAIEmbeddingFunction(
        api_key=os.environ["OPENAI_API_KEY"],
        model_name="text-embedding-3-large"
    )
    client_openai = OpenAI()
    
    coll_auth = client_chroma.get_collection("authoritative", embedding_function=embedder)
    coll_draft = client_chroma.get_collection("drafting", embedding_function=embedder)
    
    # Step 1: Get or detect opportunity
    available_opps = get_available_opportunities(client_chroma)
    print("Available opportunities:")
    for i, opp in enumerate(available_opps, 1):
        print(f"  {i}. {opp}")
    
    query = input("\nEnter opportunity name or query: ").strip()
    
    # Auto-detect opportunity
    detected_opp = detect_opportunity_from_query(query, available_opps)
    
    if detected_opp:
        print(f"✓ Detected opportunity: {detected_opp}\n")
        opportunity = detected_opp
    else:
        print("Could not auto-detect opportunity.")
        choice = input("Select opportunity number: ").strip()
        opportunity = available_opps[int(choice) - 1]
    
    # Step 2: Choose section type
    section_types = ['technical', 'management', 'past_performance', 'executive_summary']
    print("\nSection types:")
    for i, st in enumerate(section_types, 1):
        print(f"  {i}. {st.replace('_', ' ').title()}")
    
    section_choice = input("\nSelect section (1-4): ").strip()
    section_type = section_types[int(section_choice) - 1]
    
    print(f"\n📝 Generating {section_type.replace('_', ' ').title()} for {opportunity}...\n")
    
    # Step 3: Retrieve requirements
    print("→ Retrieving government requirements...")
    req_docs, req_metas = query_requirements(coll_auth, opportunity, section_type)
    print(f"  Found {len(req_docs)} requirement chunks")
    
    # Step 4: Retrieve internal best practices (optional)
    print("→ Retrieving internal best practices...")
    try:
        past_docs, past_metas = query_past_wins(coll_draft, section_type)
        print(f"  Found {len(past_docs)} past practice chunks")
    except Exception as e:
        past_docs, past_metas = [], []
        print(f"  No internal content available ({e})")
    
    # Step 5: Build context
    print("→ Building context...")
    context = build_section_context(req_docs, req_metas, past_docs, past_metas)
    
    # Step 6: Generate draft
    print("→ Generating draft with GPT-4o-mini...")
    result = generate_section(client_openai, section_type, opportunity, context, req_metas)
    
    # Step 7: Create Word document
    output_file = Path(f"{opportunity}_{section_type}_DRAFT.docx")
    print(f"→ Creating Word document: {output_file}")
    create_docx_output(section_type, opportunity, result['content'], result['citations'], output_file)
    
    print(f"\n✅ Draft complete!\n")
    print(f"Output: {output_file}")
    print(f"Requirements addressed: {len(req_metas)}")
    print(f"Past practices used: {len(past_docs)}")
    
    print("\n" + "="*60)
    print("PREVIEW (First 500 chars):")
    print("="*60)
    print(result['content'][:500] + "...")


if __name__ == "__main__":
    main()
