import json
import boto3
import os
import urllib.parse
import io
import sys
import subprocess
import re
from typing import Dict, Any
from openai import OpenAI, AuthenticationError

# Initialize AWS clients
s3 = boto3.client('s3')
sns = boto3.client('sns')

BUCKET_NAME = os.environ.get("BUCKET_NAME", "itadvisors")
SNS_TOPIC_ARN = os.environ.get("SNS_TOPIC_ARN")
KIMI_API_KEY = os.environ.get("KIMI_API_KEY")

# Validate API key
if not KIMI_API_KEY:
    print("CRITICAL: KIMI_API_KEY not set!")
else:
    # Mask key for logging (show only first 10 chars)
    masked = KIMI_API_KEY[:10] + "..." if len(KIMI_API_KEY) > 10 else "too_short"
    print(f"KIMI_API_KEY found (masked): {masked}")
    print(f"Key length: {len(KIMI_API_KEY)} characters")

# Kimi API Client - Using moonshot.ai as requested
try:
    client = OpenAI(
        api_key=KIMI_API_KEY,
        base_url="https://api.moonshot.ai/v1",  # Using .ai as requested
        timeout=600
    )
    print("Kimi API client initialized with moonshot.ai")
except Exception as e:
    print(f"ERROR initializing Kimi client: {e}")
    client = None


# =============================================================================
# DOCUMENT PARSER - Works without PyPDF2/pdfplumber
# =============================================================================

class DocumentParser:
    """Parse documents using available methods."""
    
    @staticmethod
    def parse(s3_bucket: str, s3_key: str) -> str:
        """Parse document and return text content."""
        import os
        ext = os.path.splitext(s3_key.lower())[1]
        
        response = s3.get_object(Bucket=s3_bucket, Key=s3_key)
        content = response['Body'].read()
        
        print(f"Parsing {ext} file, {len(content)} bytes")
        
        if ext == '.docx':
            return DocumentParser._parse_docx(content)
        elif ext == '.pdf':
            return DocumentParser._parse_pdf(content)
        else:
            text = content.decode('utf-8', errors='ignore')
            print(f"Text file parsed: {len(text)} characters")
            return text
    
    @staticmethod
    def _parse_docx(content: bytes) -> str:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            print(f"DOCX parsed: {len(text)} characters")
            return text
        except ImportError:
            print("WARNING: python-docx not available, returning raw")
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"DOCX parse error: {e}")
            return content.decode('utf-8', errors='ignore')
    
    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """Try multiple PDF parsers."""
        # Try PyPDF2
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = "\n".join([page.extract_text() or "" for page in reader.pages])
            if text.strip():
                print(f"PyPDF2 parsed: {len(text)} characters")
                return text
        except ImportError:
            print("PyPDF2 not available")
        except Exception as e:
            print(f"PyPDF2 error: {e}")
        
        # Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = "\n".join([p.extract_text() or "" for p in pdf.pages])
                if text.strip():
                    print(f"pdfplumber parsed: {len(text)} characters")
                    return text
        except ImportError:
            print("pdfplumber not available")
        except Exception as e:
            print(f"pdfplumber error: {e}")
        
        # Fallback: try to extract text using basic methods
        print("WARNING: No PDF parser available, using fallback")
        
        # Last resort: check if PDF contains extractable text
        try:
            # PDFs often have text embedded that can be extracted with simple methods
            text = content.decode('latin-1', errors='ignore')
            # Clean up PDF artifacts
            lines = []
            for line in text.split('\n'):
                # Filter out binary garbage
                if len(line) > 10 and all(ord(c) < 128 or c in ' \t\n\r' for c in line[:100]):
                    lines.append(line)
            if lines:
                result = '\n'.join(lines[:5000])  # Limit output
                print(f"Fallback parsed: {len(result)} characters")
                return result
        except Exception as e:
            print(f"Fallback error: {e}")
        
        return "[PDF parsing failed - no suitable parser available]"


# =============================================================================
# KIMI AGENT SWARM
# =============================================================================

class KimiAgentSwarm:
    """Kimi does everything."""
    
    def __init__(self, rfp_text: str, doc_type: str):
        self.rfp_text = rfp_text
        self.doc_type = doc_type
        self.context = {}
        
        if client is None:
            raise ValueError("Kimi API client not initialized")
        
    def call_kimi(self, system: str, user: str, max_tokens: int = 32000, temp: float = 1.0) -> str:
        """Call Kimi API with detailed error handling."""
        
        if not KIMI_API_KEY:
            raise AuthenticationError("KIMI_API_KEY environment variable is not set")
        
        # Validate key format (basic check)
        if len(KIMI_API_KEY) < 20:
            raise AuthenticationError(f"KIMI_API_KEY appears invalid (length: {len(KIMI_API_KEY)})")
        
        try:
            print(f"  Calling Kimi API...")
            print(f"  Endpoint: https://api.moonshot.ai/v1")
            print(f"  Model: kimi-k2.5")
            print(f"  Max tokens: {max_tokens}")
            
            response = client.chat.completions.create(
                model="kimi-k2.5",
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user}
                ],
                temperature=1.0,  # Kimi only accepts temperature=1
                max_tokens=max_tokens
            )
            
            content = response.choices[0].message.content
            print(f"  Response received: {len(content)} characters")
            return content
            
        except AuthenticationError as e:
            print(f"  AUTHENTICATION FAILED: {e}")
            print(f"  Your API key may be invalid or expired.")
            print(f"  Get a new key from: https://platform.moonshot.cn/")
            raise
            
        except Exception as e:
            print(f"  API CALL FAILED: {type(e).__name__}: {e}")
            raise

    def agent_analyze_document(self) -> Dict:
        """Kimi analyzes the RFP."""
        print("\n" + "="*60)
        print("AGENT 1: Document Analyzer")
        print("="*60)
        
        system = """You are an expert federal RFP analyst. Extract all key information.
Output valid JSON only."""
        
        # Limit input size
        rfp_snippet = self.rfp_text[:12000] if len(self.rfp_text) > 12000 else self.rfp_text
        
        user = f"""Analyze this {self.doc_type} and extract:

{rfp_snippet}

Output JSON:
{{
    "document_info": {{
        "type": "{self.doc_type}",
        "solicitation_number": "...",
        "agency": "...",
        "title": "...",
        "due_date": "..."
    }},
    "format_requirements": {{
        "font": "Arial",
        "font_size": 10,
        "margins": "1 inch",
        "page_limits": {{"volume_1": 30}}
    }},
    "scope_summary": "...",
    "evaluation_criteria": [],
    "key_deliverables": [],
    "special_requirements": []
}}"""
        
        try:
            output = self.call_kimi(system, user, max_tokens=16000)
            
            # Extract JSON
            json_match = re.search(r'\{.*\}', output, re.DOTALL)
            if json_match:
                analysis = json.loads(json_match.group())
                self.context['analysis'] = analysis
                print(f"  Analyzed: {analysis['document_info'].get('solicitation_number', 'Unknown')}")
                return analysis
        except Exception as e:
            print(f"  Analysis error: {e}")
        
        # Fallback
        self.context['analysis'] = {
            "document_info": {"type": self.doc_type, "solicitation_number": "UNKNOWN"},
            "format_requirements": {"font": "Arial", "font_size": 10, "margins": "1 inch"},
            "scope_summary": self.rfp_text[:500],
            "evaluation_criteria": [],
            "key_deliverables": [],
            "special_requirements": []
        }
        return self.context['analysis']

    def agent_write_all_sections(self) -> Dict[str, str]:
        """Kimi writes all proposal sections in one call."""
        print("\n" + "="*60)
        print("AGENT 2: Content Writer (All Sections)")
        print("="*60)
        
        analysis = self.context['analysis']
        
        # Truncate RFP for prompt
        rfp_snippet = self.rfp_text[:10000]
        
        system = """You are an expert proposal writer. Write a complete proposal response.
Output each section clearly marked."""
        
        user = f"""Write a complete {self.doc_type} response for:

AGENCY: {analysis['document_info'].get('agency', 'Federal Agency')}
SOLICITATION: {analysis['document_info'].get('solicitation_number', 'Unknown')}
SCOPE: {analysis.get('scope_summary', '')[:800]}

RFP CONTENT:
{rfp_snippet}

Write these sections with clear headers:

=== EXECUTIVE SUMMARY ===
[1-2 paragraphs]

=== TECHNICAL APPROACH ===
[Detailed technical content with methodology, implementation plan, security]

=== MANAGEMENT PLAN ===
[Project organization, schedule, quality control]

=== PAST PERFORMANCE ===
[2 case studies with metrics]

=== PRICING ===
[Cost breakdown with labor categories and totals]

Be specific and professional. Use bullet points where appropriate."""
        
        output = self.call_kimi(system, user, max_tokens=32000)
        
        # Parse sections
        sections = {}
        
        # Extract each section
        patterns = {
            'executive_summary': r'===\s*EXECUTIVE\s*SUMMARY\s*===(.*?)(?===|$)',
            'technical': r'===\s*TECHNICAL\s*APPROACH\s*===(.*?)(?===|$)',
            'management': r'===\s*MANAGEMENT\s*PLAN\s*===(.*?)(?===|$)',
            'past_performance': r'===\s*PAST\s*PERFORMANCE\s*===(.*?)(?===|$)',
            'pricing': r'===\s*PRICING\s*===(.*?)(?===|$)',
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, output, re.DOTALL | re.IGNORECASE)
            if match:
                sections[key] = match.group(1).strip()
                print(f"  {key}: {len(sections[key])} chars")
            else:
                sections[key] = f"[{key.replace('_', ' ').title()} section]"
                print(f"  {key}: not found, using placeholder")
        
        self.context['sections'] = sections
        return sections

    def execute(self) -> Dict[str, Any]:
        """Run all agents."""
        print(f"\n{'='*70}")
        print("STARTING KIMI AGENT SWARM")
        print(f"{'='*70}")
        
        import time
        start = time.time()
        
        self.agent_analyze_document()
        sections = self.agent_write_all_sections()
        
        elapsed = time.time() - start
        print(f"\n{'='*70}")
        print(f"COMPLETE in {elapsed:.1f}s")
        print(f"{'='*70}")
        
        return {
            'analysis': self.context['analysis'],
            'sections': sections
        }
    
    def _generate_document_code(self, sections: Dict) -> str:
        """Generate Python code to create Word document."""
        analysis = self.context['analysis']
        fmt = analysis.get('format_requirements', {})
        
        # Safely get values with defaults
        doc_info = analysis.get('document_info', {})
        solicitation = doc_info.get('solicitation_number') or 'UNKNOWN'
        font = fmt.get('font') or 'Arial'
        font_size = fmt.get('font_size') or 10
        
        # Save sections to JSON file for safe loading
        import json
        sections_file = '/tmp/sections.json'
        with open(sections_file, 'w', encoding='utf-8') as f:
            json.dump(sections, f, ensure_ascii=False)
        print(f"  Saved sections to {sections_file}")
        
        # Build code that loads sections from file
        code = '''import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    # Load sections from JSON
    with open('/tmp/sections.json', 'r', encoding='utf-8') as f:
        sections = json.load(f)
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = "''' + font + '''"
    style.font.size = Pt(''' + str(font_size) + ''')
    
    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("''' + self.doc_type + ''' RESPONSE")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 79, 114)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Solicitation: ''' + solicitation + '''")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(sections.get('executive_summary', 'Not available'))
    
    doc.add_page_break()
    
    # Technical Approach
    doc.add_heading("2. Technical Approach", level=1)
    doc.add_paragraph(sections.get('technical', 'Not available'))
    
    doc.add_page_break()
    
    # Management Plan
    doc.add_heading("3. Management Plan", level=1)
    doc.add_paragraph(sections.get('management', 'Not available'))
    
    doc.add_page_break()
    
    # Past Performance
    doc.add_heading("4. Past Performance", level=1)
    doc.add_paragraph(sections.get('past_performance', 'Not available'))
    
    doc.add_page_break()
    
    # Pricing
    doc.add_heading("5. Pricing", level=1)
    doc.add_paragraph(sections.get('pricing', 'Not available'))
    
    # Save
    doc.save("/tmp/output.docx")
    print("Document saved to /tmp/output.docx")

if __name__ == "__main__":
    main()
'''
        print(f"  Generated code: {len(code)} characters")
        return code


# =============================================================================
# DOCUMENT EXECUTOR
# =============================================================================

class DocumentExecutor:
    """Execute generated code."""
    
    @staticmethod
    def execute_code(sections: Dict[str, str], doc_type: str, solicitation: str, font: str = 'Arial', font_size: int = 10) -> bytes:
        print("\n" + "="*60)
        print("GENERATING DOCUMENT DIRECTLY")
        print("="*60)
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Set default font
            style = doc.styles['Normal']
            style.font.name = font
            style.font.size = Pt(font_size)
            
            # Cover page
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{doc_type} RESPONSE")
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 79, 114)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Solicitation: {solicitation}")
            run.font.size = Pt(12)
            
            doc.add_page_break()
            
            # Executive Summary
            doc.add_heading("1. Executive Summary", level=1)
            doc.add_paragraph(sections.get('executive_summary', 'Not available'))
            
            doc.add_page_break()
            
            # Technical Approach
            doc.add_heading("2. Technical Approach", level=1)
            doc.add_paragraph(sections.get('technical', 'Not available'))
            
            doc.add_page_break()
            
            # Management Plan
            doc.add_heading("3. Management Plan", level=1)
            doc.add_paragraph(sections.get('management', 'Not available'))
            
            doc.add_page_break()
            
            # Past Performance
            doc.add_heading("4. Past Performance", level=1)
            doc.add_paragraph(sections.get('past_performance', 'Not available'))
            
            doc.add_page_break()
            
            # Pricing
            doc.add_heading("5. Pricing", level=1)
            doc.add_paragraph(sections.get('pricing', 'Not available'))
            
            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            print(f"  Document generated: {output.tell()} bytes")
            return output.getvalue()
            
        except Exception as e:
            print(f"  Document generation error: {e}")
            import traceback
            print(traceback.format_exc())
        
        # Fallback
        return DocumentExecutor._fallback()
    
    @staticmethod
    def _fallback() -> bytes:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.left_margin = Inches(1)
        
        doc.add_heading("PROPOSAL RESPONSE", 0)
        doc.add_paragraph("Document generated with fallback formatting.")
        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()


# =============================================================================
# LAMBDA HANDLER
# =============================================================================

def lambda_handler(event, context):
    """Main entry point."""
    
    # Pre-flight checks
    print(f"\n{'='*70}")
    print("LAMBDA HANDLER STARTED")
    print(f"{'='*70}")
    
    if not KIMI_API_KEY:
        msg = "ERROR: KIMI_API_KEY environment variable not set!"
        print(msg)
        if SNS_TOPIC_ARN:
            sns.publish(TopicArn=SNS_TOPIC_ARN, Subject="Config Error", Message=msg)
        return {"statusCode": 500, "error": msg}
    
    if client is None:
        msg = "ERROR: Kimi client failed to initialize!"
        print(msg)
        if SNS_TOPIC_ARN:
            sns.publish(TopicArn=SNS_TOPIC_ARN, Subject="Config Error", Message=msg)
        return {"statusCode": 500, "error": msg}
    
    try:
        # Get file info
        record = event['Records'][0]
        source_key = record['s3']['object']['key']
        source_key = urllib.parse.unquote_plus(source_key)
        file_name = os.path.basename(source_key)
        
        print(f"Processing: {source_key}")
        
        # Determine type
        if "initialrfi" in source_key.lower():
            doc_type, dest_folder = "RFI", "draftrfi"
        elif "initialrfp" in source_key.lower():
            doc_type, dest_folder = "RFP", "draftrfp"
        else:
            doc_type, dest_folder = "RFP", "drafts"
        
        # Parse document
        rfp_text = DocumentParser.parse(BUCKET_NAME, source_key)
        
        # Run Kimi
        swarm = KimiAgentSwarm(rfp_text, doc_type)
        result = swarm.execute()
        
        # Create document
        analysis = result['analysis']
        doc_info = analysis.get('document_info', {})
        fmt = analysis.get('format_requirements', {})
        solicitation = doc_info.get('solicitation_number') or 'UNKNOWN'
        font = fmt.get('font') or 'Arial'
        font_size = fmt.get('font_size') or 10
        
        docx_bytes = DocumentExecutor.execute_code(
            result['sections'], 
            doc_type, 
            solicitation,
            font,
            font_size
        )
        
        # Save to S3
        base_name = file_name.rsplit('.', 1)[0]
        dest_key = f"{dest_folder}/{base_name}_Draft.docx"
        
        s3.put_object(
            Bucket=BUCKET_NAME,
            Key=dest_key,
            Body=docx_bytes,
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        print(f"Saved: {dest_key}")
        
        # Notify
        if SNS_TOPIC_ARN:
            sns.publish(
                TopicArn=SNS_TOPIC_ARN,
                Subject=f"Kimi {doc_type} Ready - {base_name}",
                Message=f"Document: s3://{BUCKET_NAME}/{dest_key}\nSize: {len(docx_bytes)} bytes"
            )
        
        return {
            "statusCode": 200,
            "status": "success",
            "document": dest_key,
            "size": len(docx_bytes)
        }
        
    except AuthenticationError as e:
        msg = f"Kimi Authentication Error: {e}\n\nYour API key is invalid. Get a new one from https://platform.moonshot.cn/"
        print(msg)
        if SNS_TOPIC_ARN:
            sns.publish(TopicArn=SNS_TOPIC_ARN, Subject="Kimi Auth Error", Message=msg)
        return {"statusCode": 401, "error": msg}
        
    except Exception as e:
        msg = f"Error: {e}"
        print(msg)
        import traceback
        print(traceback.format_exc())
        if SNS_TOPIC_ARN:
            sns.publish(TopicArn=SNS_TOPIC_ARN, Subject="Error", Message=msg)
        return {"statusCode": 500, "error": msg}