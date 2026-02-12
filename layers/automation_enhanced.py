import json
import boto3
import os
import urllib.parse
import io
import re
import requests
from typing import Dict, Any, List
from openai import OpenAI, AuthenticationError

# Initialize AWS clients
s3 = boto3.client('s3')
sns = boto3.client('sns')

# Environment variables
BUCKET_NAME = os.environ.get("BUCKET_NAME", "itadvisors")
SNS_TOPIC_ARN = os.environ.get("SNS_TOPIC_ARN")
KIMI_API_KEY = os.environ.get("KIMI_API_KEY")
RAG_API_URL = os.environ.get("RAG_API_URL")  # Your RAG endpoint
TAVILY_API_KEY = os.environ.get("TAVILY_API_KEY")
BRAVE_API_KEY = os.environ.get("BRAVE_API_KEY")

# Validate API keys
if not KIMI_API_KEY:
    print("CRITICAL: KIMI_API_KEY not set!")
else:
    masked = KIMI_API_KEY[:10] + "..." if len(KIMI_API_KEY) > 10 else "too_short"
    print(f"KIMI_API_KEY found (masked): {masked}")

# Kimi API Client
try:
    client = OpenAI(
        api_key=KIMI_API_KEY,
        base_url="https://api.moonshot.ai/v1",
        timeout=600
    )
    print("Kimi API client initialized")
except Exception as e:
    print(f"ERROR initializing Kimi client: {e}")
    client = None


# =============================================================================
# DOCUMENT PARSER
# =============================================================================

class DocumentParser:
    """Parse documents using available methods."""
    
    @staticmethod
    def parse(s3_bucket: str, s3_key: str) -> str:
        """Parse document and return text content."""
        ext = os.path.splitext(s3_key.lower())[1]
        
        response = s3.get_object(Bucket=s3_bucket, Key=s3_key)
        content = response['Body'].read()
        
        print(f"Parsing {ext} file, {len(content)} bytes")
        
        if ext == '.docx':
            return DocumentParser._parse_docx(content)
        elif ext == '.pdf':
            return DocumentParser._parse_pdf(content)
        else:
            text = content.decode('utf-8', errors='ignore')
            print(f"Text file parsed: {len(text)} characters")
            return text
    
    @staticmethod
    def _parse_docx(content: bytes) -> str:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            print(f"DOCX parsed: {len(text)} characters")
            return text
        except Exception as e:
            print(f"DOCX parse error: {e}")
            return content.decode('utf-8', errors='ignore')
    
    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """Parse PDF using PyPDF2."""
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = "\n".join([page.extract_text() or "" for page in reader.pages])
            if text.strip():
                print(f"PyPDF2 parsed: {len(text)} characters")
                return text
        except Exception as e:
            print(f"PyPDF2 error: {e}")
        
        return "[PDF parsing failed]"


# =============================================================================
# CONTEXT ENRICHMENT APIS
# =============================================================================

class RAGClient:
    """Client for company RAG API."""
    
    @staticmethod
    def query(query: str) -> Dict[str, Any]:
        """Query RAG API for company information."""
        if not RAG_API_URL:
            print("RAG_API_URL not configured")
            return {"error": "RAG API not configured"}
        
        try:
            print(f"  Querying RAG API: {query[:100]}...")
            response = requests.post(
                RAG_API_URL,
                json={"query": query},
                timeout=30
            )
            response.raise_for_status()
            data = response.json()
            print(f"  RAG API response: {len(str(data))} characters")
            return data
        except Exception as e:
            print(f"  RAG API error: {e}")
            return {"error": str(e)}
    
    @staticmethod
    def get_company_context(rfp_requirements: str) -> str:
        """Get relevant company information based on RFP requirements."""
        queries = [
            "past performance projects similar to " + rfp_requirements[:200],
            "company certifications and capabilities",
            "team members with relevant experience",
            "company differentiators and strengths"
        ]
        
        context_parts = []
        for query in queries:
            result = RAGClient.query(query)
            if "error" not in result and result.get("results"):
                context_parts.append(str(result["results"]))
        
        return "\n\n".join(context_parts) if context_parts else ""


class TavilyClient:
    """Client for Tavily search API."""
    
    @staticmethod
    def search(query: str, max_results: int = 5) -> List[Dict]:
        """Search using Tavily API."""
        if not TAVILY_API_KEY:
            print("TAVILY_API_KEY not configured")
            return []
        
        try:
            print(f"  Tavily search: {query[:100]}...")
            response = requests.post(
                "https://api.tavily.com/search",
                json={
                    "api_key": TAVILY_API_KEY,
                    "query": query,
                    "max_results": max_results,
                    "include_answer": True
                },
                timeout=30
            )
            response.raise_for_status()
            data = response.json()
            print(f"  Tavily found {len(data.get('results', []))} results")
            return data.get('results', [])
        except Exception as e:
            print(f"  Tavily error: {e}")
            return []
    
    @staticmethod
    def get_research_context(rfp_topic: str, technology: str) -> str:
        """Research RFP requirements using Tavily."""
        queries = [
            f"{rfp_topic} {technology} best practices 2026",
            f"{technology} implementation challenges and solutions",
            f"federal contracting {rfp_topic} requirements"
        ]
        
        context_parts = []
        for query in queries:
            results = TavilyClient.search(query, max_results=3)
            for result in results:
                context_parts.append(f"### {result.get('title', '')}\n{result.get('content', '')}")
        
        return "\n\n".join(context_parts[:10])  # Limit to top 10 snippets


class BraveClient:
    """Client for Brave Search API."""
    
    @staticmethod
    def search(query: str, count: int = 5) -> List[Dict]:
        """Search using Brave API."""
        if not BRAVE_API_KEY:
            print("BRAVE_API_KEY not configured")
            return []
        
        try:
            print(f"  Brave search: {query[:100]}...")
            response = requests.get(
                "https://api.search.brave.com/res/v1/web/search",
                headers={"X-Subscription-Token": BRAVE_API_KEY},
                params={"q": query, "count": count},
                timeout=30
            )
            response.raise_for_status()
            data = response.json()
            results = data.get('web', {}).get('results', [])
            print(f"  Brave found {len(results)} results")
            return results
        except Exception as e:
            print(f"  Brave error: {e}")
            return []
    
    @staticmethod
    def get_validation_context(technology: str, agency: str) -> str:
        """Validate and enrich information using Brave."""
        queries = [
            f"{agency} {technology} requirements 2026",
            f"{technology} compliance standards federal government"
        ]
        
        context_parts = []
        for query in queries:
            results = BraveClient.search(query, count=3)
            for result in results:
                context_parts.append(f"### {result.get('title', '')}\n{result.get('description', '')}")
        
        return "\n\n".join(context_parts[:8])  # Limit to top 8 snippets


# =============================================================================
# ENHANCED KIMI AGENT SWARM WITH CONTEXT
# =============================================================================

class EnhancedKimiAgentSwarm:
    """Kimi agent with RAG + Internet context."""
    
    def __init__(self, rfp_text: str, doc_type: str):
        self.rfp_text = rfp_text
        self.doc_type = doc_type
        self.context = {}
        self.enriched_context = {}
        
        if client is None:
            raise ValueError("Kimi API client not initialized")
        
    def call_kimi(self, system: str, user: str, max_tokens: int = 32000) -> str:
        """Call Kimi API."""
        try:
            print(f"  Calling Kimi API (max_tokens: {max_tokens})...")
            
            response = client.chat.completions.create(
                model="kimi-k2.5",
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user}
                ],
                temperature=1.0,
                max_tokens=max_tokens
            )
            
            content = response.choices[0].message.content
            print(f"  Response received: {len(content)} characters")
            return content
        except Exception as e:
            print(f"  API CALL FAILED: {type(e).__name__}: {e}")
            raise

    def enrich_context(self) -> None:
        """Gather context from RAG, Tavily, and Brave APIs."""
        print("\n" + "="*60)
        print("ENRICHING CONTEXT FROM EXTERNAL SOURCES")
        print("="*60)
        
        # Extract key requirements from RFP
        rfp_snippet = self.rfp_text[:2000]
        
        # Get company information from RAG
        print("\n📚 Fetching company data from RAG...")
        rag_context = RAGClient.get_company_context(rfp_snippet)
        self.enriched_context['rag'] = rag_context
        print(f"  RAG context: {len(rag_context)} characters")
        
        # Research using Tavily
        print("\n🔍 Researching requirements with Tavily...")
        tavily_context = TavilyClient.get_research_context(
            self.doc_type,
            rfp_snippet[:100]
        )
        self.enriched_context['tavily'] = tavily_context
        print(f"  Tavily context: {len(tavily_context)} characters")
        
        # Validate with Brave
        print("\n🌐 Validating with Brave Search...")
        brave_context = BraveClient.get_validation_context(
            rfp_snippet[:50],
            "Federal Government"
        )
        self.enriched_context['brave'] = brave_context
        print(f"  Brave context: {len(brave_context)} characters")

    def agent_analyze_document(self) -> Dict:
        """Kimi analyzes the RFP with enriched context."""
        print("\n" + "="*60)
        print("AGENT 1: Document Analyzer (with enriched context)")
        print("="*60)
        
        system = """You are an expert federal RFP analyst. Extract all key information.
Output valid JSON only."""
        
        rfp_snippet = self.rfp_text[:12000]
        
        user = f"""Analyze this {self.doc_type} and extract key information:

{rfp_snippet}

Output JSON with this structure:
{{
    "document_info": {{"type": "{self.doc_type}", "solicitation_number": "...", "agency": "...", "title": "...", "due_date": "..."}},
    "format_requirements": {{"font": "Arial", "font_size": 10, "margins": "1 inch"}},
    "scope_summary": "...",
    "evaluation_criteria": [],
    "key_deliverables": [],
    "special_requirements": []
}}"""
        
        try:
            output = self.call_kimi(system, user, max_tokens=16000)
            json_match = re.search(r'\{.*\}', output, re.DOTALL)
            if json_match:
                analysis = json.loads(json_match.group())
                self.context['analysis'] = analysis
                return analysis
        except Exception as e:
            print(f"  Analysis error: {e}")
        
        # Fallback
        self.context['analysis'] = {
            "document_info": {"type": self.doc_type, "solicitation_number": "UNKNOWN"},
            "format_requirements": {"font": "Arial", "font_size": 10, "margins": "1 inch"},
            "scope_summary": self.rfp_text[:500],
            "evaluation_criteria": [],
            "key_deliverables": [],
            "special_requirements": []
        }
        return self.context['analysis']

    def agent_write_all_sections(self) -> Dict[str, str]:
        """Kimi writes proposal with company data + internet research."""
        print("\n" + "="*60)
        print("AGENT 2: Content Writer (with RAG + Internet context)")
        print("="*60)
        
        analysis = self.context['analysis']
        rfp_snippet = self.rfp_text[:10000]
        
        # Build enriched prompt
        rag_info = self.enriched_context.get('rag', '')[:5000]
        tavily_info = self.enriched_context.get('tavily', '')[:5000]
        brave_info = self.enriched_context.get('brave', '')[:3000]
        
        system = """You are an expert proposal writer. Use the provided company information and research to write a compelling, accurate proposal response."""
        
        user = f"""Write a complete {self.doc_type} response for:

AGENCY: {analysis['document_info'].get('agency', 'Federal Agency')}
SOLICITATION: {analysis['document_info'].get('solicitation_number', 'Unknown')}
SCOPE: {analysis.get('scope_summary', '')[:800]}

RFP CONTENT:
{rfp_snippet}

COMPANY INFORMATION (from RAG):
{rag_info}

INDUSTRY RESEARCH (from Tavily):
{tavily_info}

VALIDATION DATA (from Brave):
{brave_info}

Write these sections with clear headers:

=== EXECUTIVE SUMMARY ===
[1-2 compelling paragraphs using company strengths]

=== TECHNICAL APPROACH ===
[Detailed approach using best practices from research + company capabilities]

=== MANAGEMENT PLAN ===
[Organization, schedule, quality control with company experience]

=== PAST PERFORMANCE ===
[Use REAL company past performance from RAG data]

=== PRICING ===
[Cost breakdown with labor categories]

Be specific, use actual company data, and incorporate current industry standards."""
        
        output = self.call_kimi(system, user, max_tokens=32000)
        
        # Parse sections
        sections = {}
        patterns = {
            'executive_summary': r'===\s*EXECUTIVE\s*SUMMARY\s*===(.*?)(?===|$)',
            'technical': r'===\s*TECHNICAL\s*APPROACH\s*===(.*?)(?===|$)',
            'management': r'===\s*MANAGEMENT\s*PLAN\s*===(.*?)(?===|$)',
            'past_performance': r'===\s*PAST\s*PERFORMANCE\s*===(.*?)(?===|$)',
            'pricing': r'===\s*PRICING\s*===(.*?)(?===|$)',
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, output, re.DOTALL | re.IGNORECASE)
            if match:
                sections[key] = match.group(1).strip()
                print(f"  {key}: {len(sections[key])} chars")
            else:
                sections[key] = f"[{key.replace('_', ' ').title()} section]"
        
        self.context['sections'] = sections
        return sections

    def execute(self) -> Dict[str, Any]:
        """Run enhanced agent swarm."""
        print(f"\n{'='*70}")
        print("STARTING ENHANCED KIMI AGENT SWARM WITH CONTEXT ENRICHMENT")
        print(f"{'='*70}")
        
        import time
        start = time.time()
        
        # Step 1: Enrich context from external sources
        self.enrich_context()
        
        # Step 2: Analyze document
        self.agent_analyze_document()
        
        # Step 3: Write sections with enriched context
        sections = self.agent_write_all_sections()
        
        elapsed = time.time() - start
        print(f"\n{'='*70}")
        print(f"COMPLETE in {elapsed:.1f}s")
        print(f"{'='*70}")
        
        return {
            'analysis': self.context['analysis'],
            'sections': sections,
            'enriched_context': self.enriched_context
        }


# =============================================================================
# DOCUMENT EXECUTOR
# =============================================================================

class DocumentExecutor:
    """Generate Word document."""
    
    @staticmethod
    def execute_code(sections: Dict[str, str], doc_type: str, solicitation: str, font: str = 'Arial', font_size: int = 10) -> bytes:
        print("\n" + "="*60)
        print("GENERATING DOCUMENT")
        print("="*60)
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Set default font
            style = doc.styles['Normal']
            style.font.name = font
            style.font.size = Pt(font_size)
            
            # Cover page
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{doc_type} RESPONSE")
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 79, 114)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Solicitation: {solicitation}")
            run.font.size = Pt(12)
            
            doc.add_page_break()
            
            # Add all sections
            for idx, (key, title) in enumerate([
                ('executive_summary', '1. Executive Summary'),
                ('technical', '2. Technical Approach'),
                ('management', '3. Management Plan'),
                ('past_performance', '4. Past Performance'),
                ('pricing', '5. Pricing')
            ], 1):
                doc.add_heading(title, level=1)
                doc.add_paragraph(sections.get(key, 'Not available'))
                if idx < 5:
                    doc.add_page_break()
            
            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            print(f"  Document generated: {output.tell()} bytes")
            return output.getvalue()
            
        except Exception as e:
            print(f"  Document generation error: {e}")
            import traceback
            print(traceback.format_exc())
            return DocumentExecutor._fallback()
    
    @staticmethod
    def _fallback() -> bytes:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.left_margin = Inches(1)
        
        doc.add_heading("PROPOSAL RESPONSE", 0)
        doc.add_paragraph("Document generated with fallback formatting.")
        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()


# =============================================================================
# LAMBDA HANDLER
# =============================================================================

def lambda_handler(event, context):
    """Enhanced Lambda handler with RAG + Internet context."""
    
    print(f"\n{'='*70}")
    print("ENHANCED LAMBDA HANDLER STARTED")
    print(f"{'='*70}")
    
    # Validate configuration
    if not KIMI_API_KEY:
        return {"statusCode": 500, "error": "KIMI_API_KEY not set"}
    if client is None:
        return {"statusCode": 500, "error": "Kimi client failed to initialize"}
    
    try:
        # Get file info from S3 event
        record = event['Records'][0]
        source_key = record['s3']['object']['key']
        source_key = urllib.parse.unquote_plus(source_key)
        file_name = os.path.basename(source_key)
        
        print(f"Processing: {source_key}")
        
        # Determine document type
        if "initialrfi" in source_key.lower():
            doc_type, dest_folder = "RFI", "draftrfi"
        elif "initialrfp" in source_key.lower():
            doc_type, dest_folder = "RFP", "draftrfp"
        else:
            doc_type, dest_folder = "RFP", "drafts"
        
        # Parse document
        rfp_text = DocumentParser.parse(BUCKET_NAME, source_key)
        
        # Run enhanced agent swarm with context enrichment
        swarm = EnhancedKimiAgentSwarm(rfp_text, doc_type)
        result = swarm.execute()
        
        # Create document
        analysis = result['analysis']
        doc_info = analysis.get('document_info', {})
        fmt = analysis.get('format_requirements', {})
        solicitation = doc_info.get('solicitation_number') or 'UNKNOWN'
        font = fmt.get('font') or 'Arial'
        font_size = fmt.get('font_size') or 10
        
        docx_bytes = DocumentExecutor.execute_code(
            result['sections'], 
            doc_type, 
            solicitation,
            font,
            font_size
        )
        
        # Save to S3
        base_name = file_name.rsplit('.', 1)[0]
        dest_key = f"{dest_folder}/{base_name}_Draft.docx"
        
        s3.put_object(
            Bucket=BUCKET_NAME,
            Key=dest_key,
            Body=docx_bytes,
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        print(f"✓ Saved: {dest_key}")
        
        # Notify
        if SNS_TOPIC_ARN:
            sns.publish(
                TopicArn=SNS_TOPIC_ARN,
                Subject=f"Kimi {doc_type} Ready - {base_name}",
                Message=f"""Document: s3://{BUCKET_NAME}/{dest_key}
Size: {len(docx_bytes)} bytes
Context Sources:
- RAG: {len(result.get('enriched_context', {}).get('rag', ''))} chars
- Tavily: {len(result.get('enriched_context', {}).get('tavily', ''))} chars  
- Brave: {len(result.get('enriched_context', {}).get('brave', ''))} chars"""
            )
        
        return {
            "statusCode": 200,
            "status": "success",
            "document": dest_key,
            "size": len(docx_bytes),
            "context_sources": {
                "rag": len(result.get('enriched_context', {}).get('rag', '')),
                "tavily": len(result.get('enriched_context', {}).get('tavily', '')),
                "brave": len(result.get('enriched_context', {}).get('brave', ''))
            }
        }
        
    except Exception as e:
        msg = f"Error: {e}"
        print(msg)
        import traceback
        print(traceback.format_exc())
        
        if SNS_TOPIC_ARN:
            sns.publish(TopicArn=SNS_TOPIC_ARN, Subject="Error", Message=msg)
        
        return {"statusCode": 500, "error": msg}
