"""
Analyze Sam's required proposal format from reference documents.
"""

import fitz  # PyMuPDF
from docx import Document
from pathlib import Path


def analyze_pdf_format(pdf_path):
    """Analyze PDF to extract formatting details."""
    print(f"\n{'='*70}")
    print(f"ANALYZING PDF: {Path(pdf_path).name}")
    print(f"{'='*70}\n")
    
    doc = fitz.open(pdf_path)
    
    print(f"📄 DOCUMENT INFO:")
    print(f"   Total pages: {len(doc)}")
    print(f"   Title: {doc.metadata.get('title', 'N/A')}")
    print(f"   Author: {doc.metadata.get('author', 'N/A')}")
    
    # Analyze first 3 pages in detail
    for page_num in range(min(3, len(doc))):
        page = doc[page_num]
        print(f"\n📖 PAGE {page_num + 1}:")
        print(f"   Size: {page.rect.width:.1f} x {page.rect.height:.1f} points")
        
        # Extract text blocks with formatting
        blocks = page.get_text("dict")["blocks"]
        
        print(f"   Text blocks: {len(blocks)}")
        
        for i, block in enumerate(blocks[:20]):  # First 20 blocks
            if block.get("type") == 0:  # Text block
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if text:
                            font = span.get("font", "Unknown")
                            size = span.get("size", 0)
                            color = span.get("color", 0)
                            
                            # Convert color from int to RGB
                            r = (color >> 16) & 0xFF
                            g = (color >> 8) & 0xFF
                            b = color & 0xFF
                            
                            print(f"      • {text[:60]}")
                            print(f"        Font: {font}, Size: {size:.1f}pt, Color: RGB({r},{g},{b})")
                            break  # Just first span of each line
                    break  # Just first line per block for overview
    
    # Extract full text from first page
    print(f"\n📝 FIRST PAGE FULL TEXT:")
    first_page_text = doc[0].get_text()
    lines = first_page_text.split('\n')[:30]
    for i, line in enumerate(lines, 1):
        if line.strip():
            print(f"   {i}. {line.strip()}")
    
    doc.close()


def analyze_docx_format(docx_path):
    """Analyze DOCX to extract detailed formatting."""
    print(f"\n{'='*70}")
    print(f"ANALYZING DOCX: {Path(docx_path).name}")
    print(f"{'='*70}\n")
    
    doc = Document(docx_path)
    
    # Margins
    section = doc.sections[0]
    print(f"📏 MARGINS:")
    print(f"   Top: {section.top_margin.inches:.2f}\"")
    print(f"   Bottom: {section.bottom_margin.inches:.2f}\"")
    print(f"   Left: {section.left_margin.inches:.2f}\"")
    print(f"   Right: {section.right_margin.inches:.2f}\"")
    
    # Headers/Footers
    print(f"\n📄 HEADERS & FOOTERS:")
    if section.header.paragraphs:
        header_text = " ".join([p.text for p in section.header.paragraphs if p.text.strip()])
        if header_text:
            print(f"   Header: {header_text[:80]}")
    
    if section.footer.paragraphs:
        footer_text = " ".join([p.text for p in section.footer.paragraphs if p.text.strip()])
        if footer_text:
            print(f"   Footer: {footer_text[:80]}")
    
    # First 50 paragraphs structure
    print(f"\n🏗️ DOCUMENT STRUCTURE (First 50 paragraphs):")
    for i, para in enumerate(doc.paragraphs[:50], 1):
        text = para.text.strip()
        if not text:
            print(f"   {i}. [BLANK]")
            continue
        
        style = para.style.name if para.style else "None"
        
        # Get formatting details
        details = []
        if para.runs:
            run = para.runs[0]
            if run.font.name:
                details.append(f"Font: {run.font.name}")
            if run.font.size:
                details.append(f"{run.font.size.pt}pt")
            if run.font.bold:
                details.append("Bold")
            if run.font.italic:
                details.append("Italic")
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                details.append(f"RGB({rgb[0]},{rgb[1]},{rgb[2]})")
        
        if para.alignment:
            align_map = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY"}
            details.append(align_map.get(para.alignment, ""))
        
        details_str = " | ".join(details) if details else ""
        text_preview = text[:50] + "..." if len(text) > 50 else text
        
        print(f"   {i}. [{style}] {text_preview}")
        if details_str:
            print(f"       {details_str}")
    
    # Tables
    if doc.tables:
        print(f"\n📊 TABLES: {len(doc.tables)}")
        for i, table in enumerate(doc.tables[:3], 1):
            print(f"   Table {i}: {len(table.rows)} rows × {len(table.columns)} cols")


def main():
    workspace = Path(__file__).parent
    reference_dir = workspace / "reference"
    
    # Analyze PDF
    pdf_path = reference_dir / "OnsiteITAdvisors_Technical Proposal.pdf"
    if pdf_path.exists():
        analyze_pdf_format(str(pdf_path))
    
    # Analyze DOCX
    docx_path = reference_dir / "OnsiteITAdvisors_TSC_Proposal.docx"
    if docx_path.exists():
        analyze_docx_format(str(docx_path))
    
    print(f"\n{'='*70}")
    print("ANALYSIS COMPLETE")
    print("="*70)


if __name__ == "__main__":
    main()
