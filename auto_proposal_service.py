"""
Auto-Proposal Service
Watches Federal_Contracting folder and automatically generates complete proposals.

WORKFLOW:
1. Detects new RFP files in 01_Active_Pursuits
2. Auto-ingests to ChromaDB
3. Generates all proposal sections
4. Creates compliance matrix
5. Runs requirement tracker
6. Complete proposal package ready

USAGE:
    python auto_proposal_service.py

    Service runs continuously, watching for new files.
"""

import os
import time
import threading
from pathlib import Path
from datetime import datetime
from queue import Queue
import logging

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler


# Custom formatter that removes emojis for Windows console
class SafeFormatter(logging.Formatter):
    def format(self, record):
        # Remove emoji characters for Windows compatibility
        import re
        result = super().format(record)
        # Keep only ASCII characters for console output
        result = re.sub(r'[^\x00-\x7F]+', '', result)
        return result

# Configure logging with UTF-8 for file, safe formatter for console
file_handler = logging.FileHandler('auto_proposal_service.log', encoding='utf-8')
file_handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))

console_handler = logging.StreamHandler()
console_handler.setFormatter(SafeFormatter('%(asctime)s [%(levelname)s] %(message)s'))

logging.basicConfig(
    level=logging.INFO,
    handlers=[file_handler, console_handler]
)
logger = logging.getLogger(__name__)


class RFPFileHandler(FileSystemEventHandler):
    """Monitors Federal_Contracting folder for new RFP files."""
    
    def __init__(self, job_queue: Queue):
        self.job_queue = job_queue
        self.processed_files = set()  # Track processed files to avoid duplicates
        
        # File extensions to monitor
        self.watch_extensions = {'.pdf', '.docx', '.doc', '.xlsx', '.xls'}
        
        # Folders to watch (Government-issued documents)
        self.watch_folders = [
            '01_Government_Issued',
            'Draft_Solicitations',
            'Final_Solicitations',
            'Amendments_QA'
        ]
    
    def on_created(self, event):
        """Triggered when new file is created."""
        if event.is_directory:
            # New opportunity folder created?
            self._check_new_opportunity_folder(event.src_path)
            return
        
        file_path = Path(event.src_path)
        
        # Check if it's a document we care about
        if not self._should_process(file_path):
            return
        
        # Check if already processed
        if str(file_path) in self.processed_files:
            return
        
        logger.info(f"🔔 NEW FILE DETECTED: {file_path.name}")
        
        # Extract opportunity from path
        opportunity = self._extract_opportunity(file_path)
        
        if opportunity:
            # Add to processing queue
            job = {
                'type': 'new_rfp',
                'file_path': file_path,
                'opportunity': opportunity,
                'timestamp': datetime.now()
            }
            
            self.job_queue.put(job)
            self.processed_files.add(str(file_path))
            
            logger.info(f"[QUEUED] {opportunity}")
    
    def on_modified(self, event):
        """Triggered when file is modified (amendments, updates)."""
        if event.is_directory:
            return
        
        file_path = Path(event.src_path)
        
        # Only process amendments/updates
        if 'Amendment' in file_path.name or 'Update' in file_path.name:
            if self._should_process(file_path):
                opportunity = self._extract_opportunity(file_path)
                
                if opportunity and str(file_path) not in self.processed_files:
                    logger.info(f"[UPDATED] File: {file_path.name}")
                    
                    job = {
                        'type': 'amendment',
                        'file_path': file_path,
                        'opportunity': opportunity,
                        'timestamp': datetime.now()
                    }
                    
                    self.job_queue.put(job)
                    self.processed_files.add(str(file_path))
    
    def _should_process(self, file_path: Path) -> bool:
        """Check if file should be processed."""
        # Check extension
        if file_path.suffix.lower() not in self.watch_extensions:
            return False
        
        # Check if in correct folder structure
        path_parts = file_path.parts
        
        # Must be in 01_Active_Pursuits
        if '01_Active_Pursuits' not in path_parts:
            return False
        
        # Must be in government folder
        if not any(folder in path_parts for folder in self.watch_folders):
            return False
        
        # Ignore temp files
        if file_path.name.startswith('~') or file_path.name.startswith('.'):
            return False
        
        return True
    
    def _extract_opportunity(self, file_path: Path) -> str:
        """Extract opportunity name from file path."""
        parts = file_path.parts
        
        # Find 01_Active_Pursuits index
        try:
            idx = parts.index('01_Active_Pursuits')
            # Opportunity name is next folder
            if idx + 1 < len(parts):
                return parts[idx + 1]
        except ValueError:
            pass
        
        return None
    
    def _check_new_opportunity_folder(self, folder_path: str):
        """Check if a new opportunity folder was created."""
        folder = Path(folder_path)
        
        # Check if it's directly under 01_Active_Pursuits
        if folder.parent.name == '01_Active_Pursuits':
            opportunity = folder.name
            
            # Skip template folders (including copies)
            if 'TEMPLATE' not in opportunity.upper():
                logger.info(f"[NEW OPPORTUNITY] Folder: {opportunity}")
                
                job = {
                    'type': 'new_opportunity',
                    'opportunity': opportunity,
                    'folder_path': folder,
                    'timestamp': datetime.now()
                }
                
                self.job_queue.put(job)


class ProposalPipeline:
    """Processes jobs from queue and generates proposals."""
    
    def __init__(self, workspace_path: Path):
        self.workspace_path = workspace_path
        self.active = True
    
    def run_ingestion(self, opportunity: str = None):
        """Run rag_ingest.py to process new documents."""
        logger.info(f"[INGEST] Processing documents{' for ' + opportunity if opportunity else ''}...")
        
        try:
            import subprocess
            import sys as _sys
            result = subprocess.run(
                [_sys.executable, 'rag_ingest.py'],
                cwd=self.workspace_path,
                capture_output=True,
                text=True,
                timeout=600  # 10 minute timeout
            )
            
            if result.returncode == 0:
                logger.info("[SUCCESS] Ingestion complete")
                return True
            else:
                logger.error(f"[ERROR] Ingestion failed: {result.stderr}")
                return False
        except Exception as e:
            logger.error(f"[ERROR] Ingestion error: {e}")
            return False
    
    def generate_compliance_matrix(self, opportunity: str):
        """Generate compliance matrix for opportunity."""
        logger.info(f"[MATRIX] Generating compliance matrix for {opportunity}...")
        
        try:
            # Import and run compliance matrix generator
            import sys
            sys.path.insert(0, str(self.workspace_path))
            
            from compliance_matrix import (
                extract_requirements, categorize_requirements, 
                create_compliance_matrix
            )
            import chromadb
            from chromadb.utils import embedding_functions
            
            # Setup
            client_chroma = chromadb.PersistentClient(path=str(self.workspace_path / "chroma_db"))
            embedder = embedding_functions.OpenAIEmbeddingFunction(
                api_key=os.environ["OPENAI_API_KEY"],
                model_name="text-embedding-3-large"
            )
            coll = client_chroma.get_collection("authoritative", embedding_function=embedder)
            
            # Query for requirements
            where = {
                "$and": [
                    {"opportunity": opportunity},
                    {"stage": {"$ne": "award"}},
                    {"doc_role": {"$in": ["technical_requirements", "instructions", "general", "amendment_qa"]}}
                ]
            }
            
            query_result = coll.query(
                query_texts=["requirements specifications shall must will should deliverables"],
                n_results=50,
                where=where
            )
            
            docs = query_result["documents"][0]
            metas = query_result["metadatas"][0]
            
            # Extract requirements
            all_requirements = []
            for doc, meta in zip(docs, metas):
                requirements = extract_requirements(doc)
                for req in requirements:
                    req['section'] = meta.get('filename', 'Unknown')
                    req['page'] = meta.get('page')
                all_requirements.extend(requirements)
            
            # Create matrix
            output_file = self.workspace_path / f"{opportunity}_Compliance_Matrix.xlsx"
            create_compliance_matrix(all_requirements, opportunity, output_file)
            
            logger.info(f"[SUCCESS] Compliance matrix created: {output_file.name}")
            return True
            
        except Exception as e:
            logger.error(f"[ERROR] Compliance matrix error: {e}")
            return False
    
    def generate_all_sections(self, opportunity: str):
        """Generate complete unified proposal document with all sections."""
        
        # Default sections for federal technical proposals
        # Can be customized per opportunity by creating opportunity-specific config
        default_sections = [
            ('executive_summary', 'Executive Summary'),
            ('technical', 'Technical Approach'),
            ('management', 'Management Plan'),
            ('past_performance', 'Past Performance'),
            ('staffing', 'Staffing Plan'),
            ('quality_assurance', 'Quality Assurance'),
            ('security', 'Security and Compliance'),
            ('transition', 'Transition Plan')
        ]
        
        # Check if opportunity has custom section config
        config_file = self.workspace_path / "Federal_Contracting" / "01_Active_Pursuits" / opportunity / "section_config.json"
        
        if config_file.exists():
            try:
                import json
                with open(config_file, 'r') as f:
                    config = json.load(f)
                    # Only include sections marked as required=true
                    sections = [(s['id'], s['title']) for s in config.get('sections', []) if s.get('required', False)]
                    if not sections:
                        logger.warning(f"[WARN] No required sections in config, using defaults")
                        sections = default_sections[:4]
                    else:
                        logger.info(f"[CONFIG] Using custom sections from {config_file.name}")
            except Exception as e:
                logger.warning(f"[WARN] Could not load section config, using defaults: {e}")
                sections = default_sections[:4]  # Use first 4 as safe default
        else:
            # Use first 4 sections as default (most common in federal proposals)
            sections = default_sections[:4]
            logger.info(f"[DEFAULT] Using standard 4-section format (create section_config.json to customize)")
        
        logger.info(f"[GENERATE] Creating complete proposal for {opportunity}...")
        logger.info(f"[SECTIONS] Will generate: {', '.join([s[1] for s in sections])}")
        
        # Dictionary to store generated content
        section_contents = {}
        
        for section_id, section_title in sections:
            try:
                logger.info(f"  [SECTION] Generating {section_title}...")
                content, citations = self._generate_section_content(opportunity, section_id)
                
                if content:
                    section_contents[section_id] = {
                        'title': section_title,
                        'content': content,
                        'citations': citations
                    }
                    logger.info(f"  [OK] {section_title} complete")
                else:
                    logger.warning(f"  [WARN] {section_title} failed")
                
                # Small delay between sections
                time.sleep(2)
                
            except Exception as e:
                logger.error(f"  [ERROR] {section_title} error: {e}")
        
        # Create unified proposal document
        if section_contents:
            try:
                self._create_unified_proposal(opportunity, section_contents, sections)
                logger.info(f"[COMPLETE] Generated complete proposal with {len(section_contents)}/{len(sections)} sections")
                return True
            except Exception as e:
                logger.error(f"[ERROR] Failed to create unified document: {e}")
                return False
        else:
            logger.error(f"[FAILED] No sections generated")
            return False
    
    def _generate_section_content(self, opportunity: str, section_type: str):
        """Generate content for a single proposal section."""
        try:
            import sys
            sys.path.insert(0, str(self.workspace_path))
            
            import chromadb
            from chromadb.utils import embedding_functions
            from openai import OpenAI
            
            # Initialize ChromaDB
            client_chroma = chromadb.PersistentClient(path=str(self.workspace_path / "chroma_db"))
            embedder = embedding_functions.OpenAIEmbeddingFunction(
                api_key=os.environ["OPENAI_API_KEY"],
                model_name="text-embedding-3-large"
            )
            coll_auth = client_chroma.get_collection("authoritative", embedding_function=embedder)
            coll_draft = client_chroma.get_collection("drafting", embedding_function=embedder)
            
            # Initialize OpenAI
            client_openai = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
            
            # Import functions from proposal_generator
            from proposal_generator import (
                query_requirements, query_past_wins, 
                build_section_context, generate_section
            )
            
            # Query contexts - pass collection objects as first parameter
            req_docs, req_metas = query_requirements(coll_auth, opportunity, section_type)
            past_docs, past_metas = query_past_wins(coll_draft, section_type)
            
            # Build context
            context = build_section_context(req_docs, req_metas, past_docs, past_metas)
            
            # Generate content - pass OpenAI client as first parameter
            result = generate_section(client_openai, section_type, opportunity, context, req_metas)
            
            return result['content'], result['citations']
            
        except Exception as e:
            logger.error(f"Section generation error: {e}")
            return None, []
    
    def _create_unified_proposal(self, opportunity: str, section_contents: dict, section_order: list):
        """Create unified proposal matching Sam's required format."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn
            from datetime import datetime
            
            doc = Document()
            
            # Set margins (1 inch all sides for professional federal proposals)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            
            # Add header with document title
            header = sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Onsite IT Advisors - {opportunity.replace('_', ' ')}"
            header_run = header_para.runs[0]
            header_run.font.name = 'Times New Roman'
            header_run.font.size = Pt(10)
            header_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add footer with page numbers (right-aligned)
            footer = sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            page_num_run = footer_para.add_run()
            page_num_run.font.name = 'Calibri'
            page_num_run.font.size = Pt(12)
            page_num_run.font.bold = True
            # Add field code for page number
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            page_num_run._r.append(fldChar1)
            page_num_run._r.append(instrText)
            page_num_run._r.append(fldChar2)
            footer_para.add_run(' | ')
            
            # ===== COVER PAGE (Sam's format) =====
            doc.add_paragraph("")
            doc.add_paragraph("")
            
            # Main title - large, blue, bold, centered
            title1 = doc.add_paragraph("INFORMATION TECHNOLOGY")
            title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t1_run = title1.runs[0]
            t1_run.font.name = 'Times New Roman'
            t1_run.font.size = Pt(28)
            t1_run.font.color.rgb = RGBColor(68, 114, 196)  # Sam's blue
            t1_run.font.bold = True
            
            title2 = doc.add_paragraph("OUTSOURCING AND STAFFING")
            title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t2_run = title2.runs[0]
            t2_run.font.name = 'Times New Roman'
            t2_run.font.size = Pt(28)
            t2_run.font.color.rgb = RGBColor(68, 114, 196)
            t2_run.font.bold = True
            
            title3 = doc.add_paragraph("SERVICES")
            title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t3_run = title3.runs[0]
            t3_run.font.name = 'Times New Roman'
            t3_run.font.size = Pt(28)
            t3_run.font.color.rgb = RGBColor(68, 114, 196)
            t3_run.font.bold = True
            
            doc.add_paragraph("")
            
            # RFP number/opportunity name
            rfp_para = doc.add_paragraph(opportunity.replace('_', ' '))
            rfp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rfp_run = rfp_para.runs[0]
            rfp_run.font.name = 'Times New Roman'
            rfp_run.font.size = Pt(18)
            rfp_run.font.color.rgb = RGBColor(47, 84, 150)  # Darker blue
            rfp_run.font.bold = True
            
            doc.add_paragraph("")
            doc.add_paragraph("")
            
            # Two-column layout using table
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Left column - Submitted By
            left_cell = table.rows[0].cells[0]
            left_cell.width = Inches(3.25)
            
            p = left_cell.paragraphs[0]
            run = p.add_run("Submitted By:")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(47, 84, 150)
            run.font.bold = True
            
            p = left_cell.add_paragraph("Onsite IT Advisors")
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.bold = True
            
            p = left_cell.add_paragraph("12600 Deerfield Parkway, Suite 100")
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)
            
            p = left_cell.add_paragraph("Alpharetta, Georgia 30004")
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)
            
            p = left_cell.add_paragraph("Phone: 770-940-3445")
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)
            
            p = left_cell.add_paragraph("Email: govt@onsiteitadvisors.com")
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)
            
            p = left_cell.add_paragraph("Website: www.onsiteitadvisors.com")
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)
            
            # Right column - Submitted To
            right_cell = table.rows[0].cells[1]
            right_cell.width = Inches(3.25)
            
            p = right_cell.paragraphs[0]
            run = p.add_run("Submitted To:")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(47, 84, 150)
            run.font.bold = True
            
            # Extract client info from opportunity name if possible
            client_name = opportunity.split('_')[0].replace('_', ' ')
            p = right_cell.add_paragraph(f"{client_name} Contracting Office")
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.bold = True
            
            p = right_cell.add_paragraph(f"{opportunity.replace('_', ' ')}")
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)
            
            # Remove table borders
            for row in table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
            
            doc.add_paragraph("")
            doc.add_paragraph("")
            
            # Technical Proposal label
            tech_label = doc.add_paragraph("Technical Proposal")
            tech_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tl_run = tech_label.runs[0]
            tl_run.font.name = 'Times New Roman'
            tl_run.font.size = Pt(14)
            tl_run.font.color.rgb = RGBColor(47, 84, 150)
            tl_run.font.bold = True
            
            # Page break
            doc.add_page_break()
            
            # ===== TABLE OF CONTENTS =====
            toc_heading = doc.add_heading('Table of Contents', 1)
            toc_heading_run = toc_heading.runs[0]
            toc_heading_run.font.name = 'Times New Roman'
            toc_heading_run.font.size = Pt(18)
            toc_heading_run.font.color.rgb = RGBColor(26, 54, 93)  # Dark blue
            toc_heading_run.font.bold = True
            
            # Add instruction note
            note_para = doc.add_paragraph("(Right-click and select \"Update Field\" to refresh page numbers)")
            note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            note_run = note_para.runs[0]
            note_run.font.name = 'Times New Roman'
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = RGBColor(113, 128, 150)
            note_run.font.italic = True
            
            doc.add_paragraph("")
            
            # Add TOC entries
            for idx, (section_id, section_title) in enumerate(section_order, 1):
                if section_id in section_contents:
                    toc_para = doc.add_paragraph(f"Section {idx}: {section_title}")
                    toc_para.paragraph_format.left_indent = Inches(0)
                    toc_para.paragraph_format.space_after = Pt(6)
                    for run in toc_para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
            
            # Add Source Citations to TOC
            toc_para = doc.add_paragraph("Source Citations")
            toc_para.paragraph_format.left_indent = Inches(0)
            toc_para.paragraph_format.space_after = Pt(6)
            for run in toc_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            
            doc.add_page_break()
            
            # ===== SECTIONS (Sam's format) =====
            for idx, (section_id, section_title) in enumerate(section_order, 1):
                if section_id not in section_contents:
                    continue
                
                section_data = section_contents[section_id]
                
                # Section heading - "Section X:" format, dark blue
                section_heading = doc.add_heading(f"Section {idx}: {section_title}", 1)
                section_heading_run = section_heading.runs[0]
                section_heading_run.font.name = 'Times New Roman'
                section_heading_run.font.size = Pt(14)
                section_heading_run.font.color.rgb = RGBColor(31, 56, 100)  # Sam's dark blue
                section_heading_run.font.bold = True
                
                # Section content - justified, Times New Roman 12pt
                content_paragraphs = section_data['content'].split('\n\n')
                for para_text in content_paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    
                    # Check for sub-headings (##)
                    if para_text.startswith('###'):
                        heading_text = para_text.replace('###', '').strip()
                        # Remove any markdown formatting from headings
                        heading_text = heading_text.replace('**', '').replace('*', '')
                        h3 = doc.add_heading(heading_text, 3)
                        h3_run = h3.runs[0]
                        h3_run.font.name = 'Times New Roman'
                        h3_run.font.size = Pt(12)
                        h3_run.font.color.rgb = RGBColor(0, 0, 0)
                        h3_run.font.bold = True
                    elif para_text.startswith('##'):
                        heading_text = para_text.replace('##', '').strip()
                        # Remove any markdown formatting from headings
                        heading_text = heading_text.replace('**', '').replace('*', '')
                        h2 = doc.add_heading(heading_text, 2)
                        h2_run = h2.runs[0]
                        h2_run.font.name = 'Times New Roman'
                        h2_run.font.size = Pt(12)
                        h2_run.font.color.rgb = RGBColor(0, 0, 0)
                        h2_run.font.bold = True
                    elif para_text.startswith('-') or para_text.startswith('•'):
                        # Bullet point
                        bullet_text = para_text.lstrip('-•').strip()
                        bullet_para = doc.add_paragraph(style='List Bullet')
                        bullet_para.paragraph_format.left_indent = Inches(0.5)
                        bullet_para.paragraph_format.space_after = Pt(6)
                        bullet_para.paragraph_format.line_spacing = 1.0
                        
                        # Parse markdown formatting in bullets too
                        self._add_formatted_text(bullet_para, bullet_text)
                    else:
                        # Regular paragraph - justified, Times New Roman 12pt
                        # Parse and apply markdown formatting
                        para = doc.add_paragraph()
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        para.paragraph_format.space_after = Pt(12)
                        para.paragraph_format.line_spacing = 1.0
                        
                        # Parse markdown bold and italic
                        self._add_formatted_text(para, para_text)
                
                # Spacing before next section
                doc.add_paragraph("")
            
            # ===== SOURCE CITATIONS =====
            doc.add_page_break()
            
            appendix_heading = doc.add_heading('Source Citations', 1)
            appendix_heading_run = appendix_heading.runs[0]
            appendix_heading_run.font.name = 'Times New Roman'
            appendix_heading_run.font.size = Pt(14)
            appendix_heading_run.font.color.rgb = RGBColor(31, 56, 100)
            appendix_heading_run.font.bold = True
            
            intro_para = doc.add_paragraph(
                "This proposal was developed using information from the following source documents:"
            )
            intro_para.paragraph_format.space_after = Pt(12)
            intro_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if intro_para.runs:
                intro_run = intro_para.runs[0]
                intro_run.font.name = 'Times New Roman'
                intro_run.font.size = Pt(12)
            
            # Collect all unique citations
            all_citations = set()
            for section_data in section_contents.values():
                all_citations.update(section_data['citations'])
            
            # Add citations as numbered list
            for i, citation in enumerate(sorted(all_citations), 1):
                cite_para = doc.add_paragraph(f"{i}. {citation}")
                cite_para.paragraph_format.left_indent = Inches(0.25)
                cite_para.paragraph_format.space_after = Pt(6)
                if cite_para.runs:
                    cite_run = cite_para.runs[0]
                    cite_run.font.name = 'Times New Roman'
                    cite_run.font.size = Pt(11)
            
            # Save unified document
            output_file = self.workspace_path / f"{opportunity}_Technical_Proposal.docx"
            doc.save(str(output_file))
            
            logger.info(f"[SUCCESS] Sam-approved format proposal saved: {output_file.name}")
            
        except Exception as e:
            logger.error(f"Unified document creation error: {e}")
            raise
    
    def _add_formatted_text(self, paragraph, text):
        """
        Add text to paragraph with markdown formatting converted to Word formatting.
        Converts **bold** to bold text, *italic* to italic text.
        """
        import re
        
        # Pattern to match **bold**, *italic*, or regular text
        pattern = r'(\*\*.*?\*\*|\*.*?\*|[^\*]+|\*)'
        
        parts = re.findall(pattern, text)
        
        for part in parts:
            if not part or part == '*':
                continue
            
            # Check for bold (**text**)
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                clean_text = part[2:-2]
                run = paragraph.add_run(clean_text)
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            # Check for italic (*text*)
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                clean_text = part[1:-1]
                run = paragraph.add_run(clean_text)
                run.font.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            # Regular text
            else:
                run = paragraph.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    def process_job(self, job: dict):
        """Process a single job from queue."""
        job_type = job['type']
        opportunity = job['opportunity']
        
        logger.info(f"\n{'='*70}")
        logger.info(f"[PROCESSING] JOB: {job_type.upper()} - {opportunity}")
        logger.info(f"{'='*70}\n")
        
        try:
            # Check if this is a template-filling job
            if job_type == 'new_rfp' and 'file_path' in job:
                file_path = job['file_path']
                
                # Detect if file is a government template
                is_template = self._is_government_template(file_path)
                
                if is_template:
                    logger.info("[TEMPLATE DETECTED] Government template document detected!")
                    logger.info("[MODE] Switching to TEMPLATE FILLING mode")
                    return self._process_template_job(job)
            
            # Standard workflow (create new proposal)
            return self._process_standard_job(job)
            
        except Exception as e:
            logger.error(f"❌ Pipeline error: {e}")
    
    def _is_government_template(self, file_path: Path) -> bool:
        """
        Detect if a document is a government template that needs filling.
        """
        if file_path.suffix.lower() not in ['.docx', '.doc']:
            return False
        
        try:
            from docx import Document
            import re
            
            doc = Document(str(file_path))
            
            # Check first 50 paragraphs for template indicators
            template_indicators = [
                r'\[INSERT.*?\]',
                r'\[CONTRACTOR.*?\]',
                r'\[OFFEROR.*?\]',
                r'Offeror shall provide',
                r'Contractor shall describe',
                r'\(Please describe\)',
                r'\(Contractor response\)'
            ]
            
            match_count = 0
            for para in doc.paragraphs[:50]:
                text = para.text
                for pattern in template_indicators:
                    if re.search(pattern, text, re.IGNORECASE):
                        match_count += 1
                        break
            
            # If we find 3+ template indicators, it's likely a template
            return match_count >= 3
            
        except Exception as e:
            logger.warning(f"[WARN] Could not check if template: {e}")
            return False
    
    def _process_template_job(self, job: dict):
        """Process a template-filling job."""
        opportunity = job['opportunity']
        template_path = job['file_path']
        
        logger.info(f"[STEP 1] Ingesting RFP requirements...")
        if not self.run_ingestion(opportunity):
            logger.error("[FAILED] Ingestion failed")
            return
        
        time.sleep(3)
        
        logger.info(f"[STEP 2] Detecting fillable sections in template...")
        
        # Import template filler
        import sys
        sys.path.insert(0, str(self.workspace_path))
        from template_filler import fill_template
        
        # Fill the template
        output_file = fill_template(str(template_path), opportunity, self.workspace_path)
        
        if output_file:
            logger.info(f"\n{'='*70}")
            logger.info(f"[SUCCESS] TEMPLATE FILLED: {opportunity}")
            logger.info(f"{'='*70}")
            logger.info(f"\nOutput file:")
            logger.info(f"  [DOCX] {output_file.name}")
            logger.info(f"\n[READY] Review filled template and edit as needed!\n")
        else:
            logger.error("[FAILED] Template filling failed")
    
    def _process_standard_job(self, job: dict):
        """Process a standard new-proposal job."""
        opportunity = job['opportunity']
        
        # Step 1: Ingest documents
        if not self.run_ingestion(opportunity):
            logger.error("[FAILED] Pipeline stopped: Ingestion failed")
            return
        
        time.sleep(3)  # Allow ingestion to complete
        
        # Step 2: Generate compliance matrix
        if not self.generate_compliance_matrix(opportunity):
            logger.warning("[WARN] Compliance matrix failed, continuing...")
        
        time.sleep(2)
        
        # Step 3: Generate all proposal sections
        if not self.generate_all_sections(opportunity):
            logger.error("[FAILED] Pipeline stopped: Section generation failed")
            return
        
        # Step 4: Success summary
        logger.info(f"\n{'='*70}")
        logger.info(f"[SUCCESS] PROPOSAL PACKAGE COMPLETE: {opportunity}")
        logger.info(f"{'='*70}")
        logger.info(f"\nGenerated files:")
        logger.info(f"  [XLSX] {opportunity}_Compliance_Matrix.xlsx")
        logger.info(f"  [DOCX] {opportunity}_Technical_Proposal.docx")
        logger.info(f"\n[READY] Professional proposal document ready for review!")
        logger.info(f"        Open the DOCX file to review all sections.\n")


class AutoProposalService:
    """Main service orchestrator."""
    
    def __init__(self, workspace_path: Path, watch_path: Path):
        self.workspace_path = workspace_path
        self.watch_path = watch_path
        self.job_queue = Queue()
        self.pipeline = ProposalPipeline(workspace_path)
        self.observer = None
        self.worker_thread = None
        self.running = False
    
    def start(self):
        """Start the service."""
        logger.info("="*70)
        logger.info("[START] AUTO-PROPOSAL SERVICE STARTING")
        logger.info("="*70)
        logger.info(f"Workspace: {self.workspace_path}")
        logger.info(f"Watching: {self.watch_path}")
        logger.info("="*70)
        
        # Start file watcher
        event_handler = RFPFileHandler(self.job_queue)
        self.observer = Observer()
        self.observer.schedule(event_handler, str(self.watch_path), recursive=True)
        self.observer.start()
        
        logger.info("[ACTIVE] File watcher monitoring")
        
        # Start job processor
        self.running = True
        self.worker_thread = threading.Thread(target=self._process_jobs, daemon=True)
        self.worker_thread.start()
        
        logger.info("[ACTIVE] Job processor ready")
        logger.info("\n[READY] SERVICE READY - Watching for new RFP files...\n")
        
        try:
            # Keep service running
            while self.running:
                time.sleep(1)
        except KeyboardInterrupt:
            self.stop()
    
    def stop(self):
        """Stop the service."""
        logger.info("\n[STOP] Stopping service...")
        
        self.running = False
        
        if self.observer:
            self.observer.stop()
            self.observer.join()
        
        logger.info("[STOPPED] Service stopped")
    
    def _process_jobs(self):
        """Worker thread that processes jobs from queue."""
        while self.running:
            try:
                # Wait for job (timeout allows checking running flag)
                if not self.job_queue.empty():
                    job = self.job_queue.get(timeout=1)
                    
                    # Process job
                    self.pipeline.process_job(job)
                    
                    self.job_queue.task_done()
                else:
                    time.sleep(1)
                    
            except Exception as e:
                logger.error(f"Worker error: {e}")
                time.sleep(5)


def main():
    # Check for OpenAI API key
    if not os.getenv("OPENAI_API_KEY"):
        logger.error("[ERROR] OPENAI_API_KEY not set")
        logger.error("Set environment variable: $env:OPENAI_API_KEY='your-key'")
        return
    
    # Paths
    workspace_path = Path(__file__).parent
    watch_path = workspace_path / "Federal_Contracting" / "01_Active_Pursuits"
    
    if not watch_path.exists():
        logger.error(f"[ERROR] Watch path does not exist: {watch_path}")
        return
    
    # Create and start service
    service = AutoProposalService(workspace_path, watch_path)
    service.start()


if __name__ == "__main__":
    main()
