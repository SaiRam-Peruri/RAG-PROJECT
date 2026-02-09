"""
Template Filling Mode
For RFPs where government provides a template document that must be filled in.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
import re
import logging

logger = logging.getLogger(__name__)


def _add_formatted_text_to_paragraph(paragraph, text):
    """
    Add text to paragraph with markdown formatting converted to Word formatting.
    Converts **bold** to bold text, *italic* to italic text.
    """
    # Pattern to match **bold**, *italic*, or regular text
    pattern = r'(\*\*.*?\*\*|\*.*?\*|[^\*]+|\*)'
    
    parts = re.findall(pattern, text)
    
    for part in parts:
        if not part or part == '*':
            continue
        
        # Check for bold (**text**)
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            clean_text = part[2:-2]
            run = paragraph.add_run(clean_text)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        # Check for italic (*text*)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
            clean_text = part[1:-1]
            run = paragraph.add_run(clean_text)
            run.font.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        # Regular text
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def detect_fillable_sections(doc_path: str):
    """
    Analyze a government template to find fillable sections.
    Returns list of sections that need contractor responses.
    """
    doc = Document(doc_path)
    
    fillable_sections = []
    
    # Common indicators of fillable sections in government templates
    fill_indicators = [
        r'\[INSERT.*?\]',
        r'\[CONTRACTOR.*?\]',
        r'\[OFFEROR.*?\]',
        r'\[YOUR RESPONSE.*?\]',
        r'Offeror shall provide',
        r'Contractor shall describe',
        r'Provide a description of',
        r'___+',  # Blank lines
        r'\(Please describe\)',
        r'\(Contractor response\)',
    ]
    
    current_section = None
    section_counter = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        # Check if this is a section heading (bold, larger font, or numbered)
        is_heading = False
        if para.style and 'Heading' in para.style.name:
            is_heading = True
        elif para.runs:
            first_run = para.runs[0]
            if first_run.font.bold and first_run.font.size and first_run.font.size.pt >= 14:
                is_heading = True
        
        if is_heading:
            current_section = {
                'paragraph_index': i,
                'heading': text,
                'fill_paragraphs': []
            }
        
        # Check if this paragraph needs filling
        needs_fill = False
        for pattern in fill_indicators:
            if re.search(pattern, text, re.IGNORECASE):
                needs_fill = True
                break
        
        if needs_fill and current_section:
            current_section['fill_paragraphs'].append({
                'paragraph_index': i,
                'original_text': text,
                'style': para.style.name if para.style else 'Normal'
            })
            
            # If we haven't added this section yet, add it
            if current_section not in fillable_sections:
                fillable_sections.append(current_section)
                section_counter += 1
    
    return fillable_sections


def map_section_to_type(section_heading: str):
    """
    Map a government template section heading to our section types.
    """
    heading_lower = section_heading.lower()
    
    mapping = {
        'executive summary': 'executive_summary',
        'technical approach': 'technical',
        'technical solution': 'technical',
        'management': 'management',
        'management plan': 'management',
        'organizational': 'management',
        'past performance': 'past_performance',
        'relevant experience': 'past_performance',
        'staffing': 'staffing',
        'personnel': 'staffing',
        'key personnel': 'staffing',
        'quality': 'quality_assurance',
        'qa/qc': 'quality_assurance',
        'security': 'security',
        'compliance': 'security',
        'transition': 'transition',
        'cost': 'cost',
        'pricing': 'cost'
    }
    
    for keyword, section_type in mapping.items():
        if keyword in heading_lower:
            return section_type
    
    return 'technical'  # Default


def fill_template(template_path: str, opportunity: str, workspace_path: Path):
    """
    Fill in a government-provided template document with auto-generated content.
    
    Process:
    1. Detect fillable sections in template
    2. Generate content for each section
    3. Replace placeholder text with generated content
    4. Save as new file (preserving all government formatting)
    """
    logger.info(f"[TEMPLATE MODE] Analyzing government template: {Path(template_path).name}")
    
    # Detect fillable sections
    fillable_sections = detect_fillable_sections(template_path)
    
    if not fillable_sections:
        logger.warning("[WARN] No fillable sections detected in template!")
        logger.warning("[WARN] Template might not have standard fill indicators")
        return None
    
    logger.info(f"[DETECTED] Found {len(fillable_sections)} fillable sections:")
    for section in fillable_sections:
        logger.info(f"  • {section['heading']} ({len(section['fill_paragraphs'])} paragraphs to fill)")
    
    # Load the template document
    doc = Document(template_path)
    
    # Import generation functions
    import sys
    sys.path.insert(0, str(workspace_path))
    
    from auto_proposal_service import ProposalPipeline
    pipeline = ProposalPipeline(workspace_path)
    
    # Generate content for each fillable section
    for section in fillable_sections:
        section_type = map_section_to_type(section['heading'])
        
        logger.info(f"[FILL] Generating content for: {section['heading']} (type: {section_type})")
        
        # Generate content
        try:
            content, citations = pipeline._generate_section_content(opportunity, section_type)
            
            if not content:
                logger.warning(f"[WARN] No content generated for {section['heading']}")
                continue
            
            # Fill in the paragraphs
            for fill_para in section['fill_paragraphs']:
                para_idx = fill_para['paragraph_index']
                para = doc.paragraphs[para_idx]
                
                # Clear existing runs from paragraph
                for run in para.runs:
                    run.text = ''
                
                # Add formatted content (converting markdown to Word formatting)
                _add_formatted_text_to_paragraph(para, content)
                
                logger.info(f"  ✓ Filled paragraph at index {para_idx}")
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to fill {section['heading']}: {e}")
            continue
    
    # Save filled template
    output_file = workspace_path / f"{opportunity}_FILLED_TEMPLATE.docx"
    doc.save(str(output_file))
    
    logger.info(f"[SUCCESS] Filled template saved: {output_file.name}")
    return output_file


def fill_pdf_template(pdf_path: str, opportunity: str, workspace_path: Path):
    """
    Fill in a PDF template with form fields.
    Uses PyPDF2 to detect and fill form fields.
    """
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        logger.error("[ERROR] pypdf not installed. Run: pip install pypdf")
        return None
    
    logger.info(f"[PDF TEMPLATE MODE] Analyzing PDF template")
    
    reader = PdfReader(pdf_path)
    writer = PdfWriter()
    
    # Check if PDF has form fields
    if '/AcroForm' not in reader.trailer['/Root']:
        logger.warning("[WARN] PDF does not have fillable form fields")
        logger.info("[INFO] Will create a text overlay instead")
        return None
    
    # TODO: Implement PDF form field filling
    # This requires:
    # 1. Detect form field names
    # 2. Map field names to section types
    # 3. Generate content for each field
    # 4. Fill the fields using PdfWriter.update_page_form_field_values()
    
    logger.error("[ERROR] PDF form filling not yet implemented")
    logger.info("[INFO] Convert PDF to Word template or use Word version of template")
    return None


if __name__ == "__main__":
    # Test mode
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python template_filler.py <template.docx> <opportunity_name>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    opportunity = sys.argv[2]
    workspace = Path(__file__).parent
    
    result = fill_template(template_path, opportunity, workspace)
    
    if result:
        print(f"\n✓ Template filled successfully: {result}")
    else:
        print("\n✗ Template filling failed")
